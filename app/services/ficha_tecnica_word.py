"""
Extracción de fichas técnicas desde Word (FT *.docx) y mapeo al formato YAML del panel.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document

from app.services.ficha_tecnica import (
    FICHAS_DIR,
    FICHAS_PDF_DIR,
    _filas_tabla,
    _normalizar,
    _valor_en_filas,
    guardar_yaml_datos,
    generar_pdf_html,
    nombre_archivo_desde_titulo,
    normalizar_datos_ficha,
)

IDENTIDAD_KEYWORDS = {
    "nombre del producto",
    "cas #",
    "cas#",
    "formula molecular",
    "sinonimos",
    "nombre iupac",
}
PROPIEDADES_KEYWORDS = {
    "apariencia",
    "peso molecular",
    "concentracion",
    "rotacion",
    "perdidas",
    "punto",
    "ph",
    "solubilidad",
    "densidad",
    "viscosidad",
    "pureza",
    "humedad",
    "identificacion",
}
MICRO_KEYWORDS = {
    "mesofilos",
    "hongos",
    "levaduras",
    "e coli",
    "salmonella",
    "recuento",
    "coliformes",
    "aerobios",
}

SECTION_HEADINGS = {
    "descripcion",
    "beneficios",
    "aplicaciones",
    "propiedades",
    "recomendaciones",
    "dosificacion",
    "estabilidad",
    "almacenamiento",
    "uso",
    "usos",
    "instrucciones",
    "advertencias",
    "propiedades fisico quimicas",
    "propiedades microbiologicas",
}

_FISICAS_MAP = (
    ("apariencia", ("apariencia", "aspecto")),
    ("punto_fusion", ("punto de fusion", "punto fusion", "melting point")),
    ("indice_saponificacion", ("indice de saponificacion", "saponificacion")),
    ("ph", ("ph", "ph (1%", "ph 1%")),
    ("olor", ("olor", "odour", "odor")),
    ("sabor", ("sabor", "taste")),
    ("formula_quimica", ("formula quimica", "formula molecular", "formula")),
    ("solubilidad", ("solubilidad", "solubility")),
)


def clave_desde_nombre(nombre: str) -> str:
    return _normalizar(nombre)


def clave_desde_archivo(nombre_archivo: str) -> str:
    """FT ACIDO ASCORBICO.docx → acido ascorbico"""
    nombre = Path(nombre_archivo).stem.strip()
    if nombre.upper().startswith("FT "):
        nombre = nombre[3:].strip()
    return clave_desde_nombre(nombre)


def clasificar_tabla(tabla) -> str:
    textos = []
    for row in tabla.rows[:3]:
        for cell in row.cells[:2]:
            textos.append(_normalizar(cell.text))
    todos = " ".join(textos)
    hits_micro = sum(1 for k in MICRO_KEYWORDS if k in todos)
    hits_id = sum(1 for k in IDENTIDAD_KEYWORDS if k in todos)
    hits_prop = sum(1 for k in PROPIEDADES_KEYWORDS if k in todos)
    if hits_micro >= 1:
        return "microbiologia"
    if hits_id >= 1:
        return "identidad"
    if hits_prop >= 1:
        return "propiedades"
    return "otras"


def tabla_a_filas(tabla) -> list[list[str]]:
    filas: list[list[str]] = []
    vistas: set[tuple[str, ...]] = set()
    for row in tabla.rows:
        textos = [c.text.strip() for c in row.cells]
        clave = tuple(t[:30] for t in textos)
        if clave in vistas:
            continue
        vistas.add(clave)
        contenido = [t for t in textos if t]
        if contenido and len(contenido) >= 2:
            filas.append([textos[0].strip(), textos[-1].strip()])
        elif contenido:
            filas.append([textos[0].strip(), ""])
    return filas


def extraer_ficha(path: Path) -> dict[str, Any]:
    """Lee un .docx FT y extrae título, descripción, secciones y tablas."""
    doc = Document(str(path))

    ficha: dict[str, Any] = {
        "titulo": "",
        "descripcion": "",
        "secciones": [],
        "identidad": [],
        "propiedades": [],
        "microbiologia": [],
    }

    for tabla in doc.tables:
        tipo = clasificar_tabla(tabla)
        filas = tabla_a_filas(tabla)
        if tipo == "identidad":
            ficha["identidad"].extend(filas)
        elif tipo == "propiedades":
            ficha["propiedades"].extend(filas)
        elif tipo == "microbiologia":
            ficha["microbiologia"].extend(filas)

    seccion_actual = None
    items_actuales: list[str] = []
    titulo_extraido = False

    for para in doc.paragraphs:
        texto = para.text.strip()
        if not texto:
            continue

        texto_norm = _normalizar(texto)

        if not titulo_extraido and len(texto) < 80:
            ficha["titulo"] = texto
            titulo_extraido = True
            continue

        es_cabecera = (
            para.style.name in ("Heading 1", "Heading 2", "Heading 3")
            or texto_norm in SECTION_HEADINGS
            or (len(texto) < 60 and texto.isupper())
        )

        if es_cabecera:
            if seccion_actual and items_actuales:
                ficha["secciones"].append({"titulo": seccion_actual, "items": items_actuales})
                items_actuales = []
            seccion_actual = texto
        elif seccion_actual:
            titulo_norm = _normalizar(seccion_actual)
            if "descripci" in titulo_norm and not ficha["descripcion"]:
                ficha["descripcion"] = texto
            else:
                items_actuales.append(texto)

    if seccion_actual and items_actuales:
        ficha["secciones"].append({"titulo": seccion_actual, "items": items_actuales})

    if not ficha["descripcion"]:
        for sec in ficha["secciones"]:
            if "descripci" in _normalizar(sec["titulo"]) and sec["items"]:
                ficha["descripcion"] = sec["items"][0]
                break

    return ficha


def extraer_ficha_yaml(path: Path) -> dict[str, Any]:
    """Lee YAML del panel y lo convierte al esquema de extraer_ficha()."""
    import yaml

    with open(path, encoding="utf-8") as f:
        datos = yaml.safe_load(f) or {}

    ficha: dict[str, Any] = {
        "titulo": (datos.get("titulo") or datos.get("nombre_producto") or "").strip(),
        "descripcion": (datos.get("descripcion") or "").strip(),
        "secciones": [],
        "identidad": datos.get("identidad") or [],
        "propiedades": datos.get("propiedades") or [],
        "microbiologia": (datos.get("_coa") or {}).get("parametros") or datos.get("microbiologia") or [],
    }

    beneficios = []
    for linea in datos.get("propiedades_lista") or []:
        titulo_prop, _, desc = str(linea).partition("|")
        titulo_prop = titulo_prop.strip()
        desc = desc.strip()
        beneficios.append(f"{titulo_prop}: {desc}" if desc else titulo_prop)
    if beneficios:
        ficha["secciones"].append({"titulo": "BENEFICIOS", "items": beneficios})

    aplicaciones = [str(a).strip() for a in (datos.get("aplicaciones") or []) if str(a).strip()]
    if aplicaciones:
        ficha["secciones"].append({"titulo": "APLICACIONES", "items": aplicaciones})

    modo_uso = (datos.get("modo_uso") or "").strip()
    if modo_uso:
        ficha["secciones"].append({"titulo": "MODO DE USO", "items": [modo_uso]})

    recomendaciones = (datos.get("recomendaciones") or "").strip()
    if recomendaciones:
        ficha["secciones"].append({"titulo": "RECOMENDACIONES", "items": [recomendaciones]})

    return ficha


def _items_seccion(secciones: list, *claves: str) -> list[str]:
    claves_n = {_normalizar(c) for c in claves}
    out: list[str] = []
    for sec in secciones or []:
        titulo = _normalizar(sec.get("titulo") or "")
        if any(k in titulo for k in claves_n) or titulo in claves_n:
            for item in sec.get("items") or []:
                t = str(item).strip()
                if t:
                    out.append(t)
    return out


def _partir_item_app(item: str) -> list[str]:
    """Si una aplicación viene como viñetas en un solo párrafo, intenta separar."""
    t = item.strip()
    if "\n" in t:
        return [ln.strip(" •-\t") for ln in t.splitlines() if ln.strip(" •-\t")]
    if re.search(r"[;•]", t) and len(t) > 80:
        partes = re.split(r"[;•]+", t)
        limpios = [p.strip(" -") for p in partes if p.strip(" -")]
        if len(limpios) >= 2:
            return limpios
    return [t]


def word_ficha_a_datos(ficha: dict[str, Any], *, titulo_fallback: str = "") -> dict[str, Any]:
    """Convierte la salida de extraer_ficha() al dict del formulario FT / YAML."""
    identidad = _filas_tabla(ficha.get("identidad"))
    propiedades = _filas_tabla(ficha.get("propiedades"))
    microbiologia = _filas_tabla(ficha.get("microbiologia"))
    secciones = ficha.get("secciones") or []

    nombre = (ficha.get("titulo") or "").strip()
    if not nombre:
        nombre = _valor_en_filas(identidad, "nombre del producto") or titulo_fallback
    nombre = nombre.strip()

    def _limpio(s: str) -> str:
        # Quitar BOM / zero-width / NBSP residuales de Word
        return re.sub(r"[\u200b\u200c\u200d\ufeff\xa0]", "", (s or "")).strip()

    cas = _limpio(_valor_en_filas(identidad, "cas", "cas #", "cas#"))
    cas = re.sub(r"^(CAS\s*#?\s*)", "", cas, flags=re.I).strip() or cas
    sinonimos = _limpio(_valor_en_filas(identidad, "sinonimos", "sinonimo"))
    ref = _limpio(
        _valor_en_filas(
            identidad,
            "referencia siigo",
            "referencia interna",
            "referencia",
            "codigo",
        )
    )
    fecha = _limpio(_valor_en_filas(identidad, "fecha de revision", "fecha revision", "fecha"))

    descripcion = _limpio(ficha.get("descripcion") or "")
    if not descripcion:
        for item in _items_seccion(secciones, "descripcion"):
            descripcion = item
            break

    cf: dict[str, str] = {
        "apariencia": "",
        "punto_fusion": "",
        "indice_saponificacion": "",
        "ph": "",
        "olor": "",
        "sabor": "",
        "formula_quimica": "",
        "solubilidad": "",
    }
    usadas: set[str] = set()
    for key, aliases in _FISICAS_MAP:
        for label, val in propiedades:
            ln = _normalizar(label)
            if any(a in ln or ln == a for a in aliases) and val.strip():
                cf[key] = val.strip()
                usadas.add(ln)
                break

    # Fórmula molecular a veces está en identidad
    if not cf["formula_quimica"]:
        fm = _valor_en_filas(identidad, "formula molecular", "formula quimica", "formula")
        if fm:
            cf["formula_quimica"] = fm

    props_extra: list[list[str]] = []
    props_lista: list[str] = []
    for label, val in propiedades:
        ln = _normalizar(label)
        if ln in usadas:
            continue
        if label.strip() or val.strip():
            props_extra.append([label.strip(), val.strip()])
            if val.strip():
                props_lista.append(f"{label.strip()}|{val.strip()}")
            else:
                props_lista.append(label.strip())

    apps_raw = _items_seccion(secciones, "aplicaciones", "aplicacion", "usos", "uso")
    aplicaciones: list[str] = []
    for item in apps_raw:
        for parte in _partir_item_app(item):
            if parte and parte not in aplicaciones:
                aplicaciones.append(parte)

    # Beneficios → propiedades_lista si no hay tablas de props útiles
    beneficios = _items_seccion(secciones, "beneficios", "propiedades")
    for b in beneficios:
        # Evitar reinyectar encabezados genéricos ya mapeados como tablas
        if _normalizar(b) in ("propiedades", "beneficios"):
            continue
        if "|" not in b and ":" in b and len(b) < 200:
            a, _, c = b.partition(":")
            linea = f"{a.strip()}|{c.strip()}"
        else:
            linea = b
        if linea not in props_lista:
            props_lista.append(linea)

    reco_items = _items_seccion(
        secciones,
        "recomendaciones",
        "estabilidad",
        "almacenamiento",
        "estabilidad y almacenamiento",
    )
    recomendaciones = "\n".join(reco_items).strip()

    modo_items = _items_seccion(secciones, "modo de uso", "dosificacion", "instrucciones")
    modo_uso = "\n".join(modo_items).strip()

    estabilidad = reco_items if reco_items else []

    datos: dict[str, Any] = {
        "titulo": nombre.upper() if nombre else "PRODUCTO",
        "nombre_producto": nombre or "PRODUCTO",
        "referencia": ref,
        "sinonimos": sinonimos,
        "cas": cas,
        "fecha_revision": fecha,
        "descripcion": descripcion,
        "caracteristicas_fisicas": cf,
        "propiedades_lista": props_lista,
        "aplicaciones": aplicaciones,
        "composicion": [],
        "recomendaciones": recomendaciones,
        "modo_uso": modo_uso,
        "identidad": identidad,
        "propiedades": props_extra,
        "microbiologia": microbiologia,
        "estabilidad": estabilidad,
        "origen_word": True,
    }
    return normalizar_datos_ficha(datos)


def listar_ft_docx(directorio: Path | None = None) -> list[Path]:
    base = directorio or FICHAS_DIR
    return sorted(
        p
        for p in base.glob("FT *.docx")
        if not p.name.startswith("~$") and p.is_file()
    )


def ruta_pdf_esperada(titulo: str) -> Path:
    nombre = nombre_archivo_desde_titulo(titulo).replace(".docx", ".pdf")
    return FICHAS_PDF_DIR / nombre


def pdf_ya_existe(titulo: str) -> bool:
    path = ruta_pdf_esperada(titulo)
    if path.is_file():
        return True
    # También en raíz histórica
    alt = FICHAS_DIR / path.name
    return alt.is_file()


def campos_vacios(datos: dict[str, Any]) -> list[str]:
    vacios: list[str] = []
    if not (datos.get("descripcion") or "").strip():
        vacios.append("descripcion")
    if not (datos.get("cas") or "").strip():
        vacios.append("cas")
    if not (datos.get("aplicaciones") or []):
        vacios.append("aplicaciones")
    cf = datos.get("caracteristicas_fisicas") or {}
    if not any((cf.get(k) or "").strip() for k in ("apariencia", "olor", "solubilidad", "ph")):
        if not (datos.get("propiedades") or []):
            vacios.append("propiedades")
    return vacios


def procesar_word_a_ficha(
    path: Path,
    *,
    guardar_yaml: bool = True,
    generar_pdf: bool = True,
    force: bool = False,
    dry_run: bool = False,
    cabezote_id: str | None = None,
) -> dict[str, Any]:
    """Extrae un Word FT, mapea al formato, guarda YAML y genera PDF HTML."""
    path = Path(path)
    if not path.is_file():
        raise FileNotFoundError(f"No existe: {path}")

    ficha = extraer_ficha(path)
    titulo_fb = clave_desde_archivo(path.name).upper()
    datos = word_ficha_a_datos(ficha, titulo_fallback=titulo_fb)
    titulo = (datos.get("titulo") or titulo_fb).strip()
    vacios = campos_vacios(datos)
    pdf_path = ruta_pdf_esperada(titulo)
    ya_hay_pdf = pdf_ya_existe(titulo)

    resultado: dict[str, Any] = {
        "ok": True,
        "archivo": path.name,
        "titulo": titulo,
        "vacios": vacios,
        "pdf_existente": ya_hay_pdf,
        "pdf": str(pdf_path) if ya_hay_pdf or generar_pdf else "",
        "yaml": "",
        "dry_run": dry_run,
        "skipped": False,
    }

    if ya_hay_pdf and not force and generar_pdf:
        resultado["skipped"] = True
        resultado["motivo"] = "pdf_existente"
        return resultado

    if dry_run:
        resultado["yaml"] = f"datos/{_normalizar(titulo).replace(' ', '_')}.yaml"
        return resultado

    if guardar_yaml:
        yaml_path = guardar_yaml_datos(datos)
        resultado["yaml"] = str(yaml_path)

    if generar_pdf:
        FICHAS_PDF_DIR.mkdir(parents=True, exist_ok=True)
        res_pdf = generar_pdf_html(datos, cabezote_id=cabezote_id, salida=pdf_path)
        resultado["pdf"] = res_pdf.get("pdf") or str(pdf_path)
        resultado["pdf_nombre"] = res_pdf.get("pdf_nombre") or pdf_path.name

    return resultado
