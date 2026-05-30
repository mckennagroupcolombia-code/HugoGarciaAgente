"""
Generación de hojas de datos de seguridad (SDS) McKenna — formato GHS Ventós (ref. SDS ELEMI).
"""

from __future__ import annotations

import os
import re
import shutil
import unicodedata
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document

from app.services.ficha_tecnica import (
    CREDS_PATH,
    DATOS_DIR,
    TDS_IMPERSONATE,
    cargar_datos_desde_archivo,
    client_email_servicio,
    exportar_pdf,
    guardar_yaml_datos,
    subir_a_drive,
)

REPO = Path(__file__).resolve().parents[2]
SDS_DIR = REPO / "fichas_word"
PLANTILLA_DEFAULT = SDS_DIR / "plantillas" / "SDS PLANTILLA.docx"
PLANTILLA_REF_PDF = SDS_DIR / "plantillas" / "SDS REFERENCIA ELEMI.pdf"
PLANTILLA_YAML = DATOS_DIR / "sds_plantilla_ejemplo.yaml"

DOCS_DRIVE_PARENT = os.getenv("DOCS_DRIVE_PARENT_ID", "1fgf9W0ifD5bjN9QSLWZxvtEWby_jT16v")
SDS_DRIVE_FOLDER = os.getenv("SDS_DRIVE_FOLDER_ID", "1lbnrVKDIH4CPL6SRWQxAx27zMZqma5IU")

# Claves YAML → placeholder plantilla sección 9
_MAPA_PROPIEDADES = {
    "aspecto": "PF_ASPECTO",
    "color": "PF_COLOR",
    "olor": "PF_OLOR",
    "umbral_olfativo": "PF_UMBRAL_OLOR",
    "ph": "PF_PH",
    "punto_fusion": "PF_FUSION",
    "punto_ebullicion": "PF_EBULICION",
    "punto_inflamacion": "PF_INFLAMACION",
    "velocidad_evaporacion": "PF_EVAPORACION",
    "inflamabilidad": "PF_INFLAMABILIDAD",
    "limite_inf_inflamabilidad": "PF_LIM_INF",
    "limite_sup_inflamabilidad": "PF_LIM_SUP",
    "presion_vapor": "PF_PRESION_VAPOR",
    "densidad_vapor": "PF_DENS_VAPOR",
    "densidad": "PF_DENSIDAD",
    "densidad_relativa": "PF_DENS_REL",
    "solubilidad_agua": "PF_SOL_AGUA",
    "solubilidad_otros": "PF_SOL_OTROS",
    "logp": "PF_LOGP",
    "autoignicion": "PF_AUTOIGNICION",
    "descomposicion": "PF_DESCOMP",
    "viscosidad_dinamica": "PF_VIS_DIN",
    "viscosidad_cinematica": "PF_VIS_CIN",
    "explosivas": "PF_EXPLOSIVAS",
    "comburentes": "PF_COMBURENTES",
}


def _normalizar(s: str) -> str:
    t = unicodedata.normalize("NFD", (s or "").lower())
    t = "".join(c for c in t if unicodedata.category(c) != "Mn")
    return re.sub(r"\s+", " ", t).strip()


def nombre_archivo_desde_titulo(titulo: str) -> str:
    t = unicodedata.normalize("NFD", titulo.upper())
    t = "".join(c for c in t if unicodedata.category(c) != "Mn")
    t = re.sub(r"[^A-Z0-9]+", "-", t).strip("-")
    return f"SDS-{t}.docx"


def plantilla_datos_ejemplo() -> dict:
    if PLANTILLA_YAML.exists():
        return cargar_datos_desde_archivo(PLANTILLA_YAML)
    return {"titulo": "NOMBRE DEL PRODUCTO", "ghs": {}}


def listar_yaml_datos() -> list[dict[str, str]]:
    if not DATOS_DIR.is_dir():
        return []
    items = []
    for p in sorted(DATOS_DIR.glob("sds_*.yaml")) + sorted(DATOS_DIR.glob("sds_*.yml")):
        if "plantilla" in p.name:
            continue
        try:
            d = cargar_datos_desde_archivo(p)
            titulo = (d.get("titulo") or p.stem).strip()
        except Exception:
            titulo = p.stem
        items.append({"id": p.stem, "archivo": p.name, "titulo": titulo})
    return items


def _propiedades_a_mapa(datos: dict) -> dict[str, str]:
    """Convierte propiedades_fisicas dict o lista [[k,v],...] a claves PF_*."""
    out: dict[str, str] = {}
    pf = datos.get("propiedades_fisicas")
    if isinstance(pf, dict):
        for k, v in pf.items():
            key = _normalizar(str(k)).replace(" ", "_")
            ph = _MAPA_PROPIEDADES.get(key)
            if ph:
                out[ph] = str(v)
    for item in datos.get("propiedades") or []:
        if isinstance(item, (list, tuple)) and len(item) >= 2:
            label, val = str(item[0]), str(item[1])
            nk = _normalizar(label)
            for yaml_key, ph in _MAPA_PROPIEDADES.items():
                if yaml_key in nk or nk.startswith(yaml_key[:4]):
                    out[ph] = val
                    break
    return out


def _reemplazar_en_doc(doc: Document, mapping: dict[str, str]) -> None:
    for key, val in mapping.items():
        token = f"{{{{{key}}}}}"
        if not val:
            continue
        for p in doc.paragraphs:
            if token in p.text:
                for run in p.runs:
                    if token in run.text:
                        run.text = run.text.replace(token, val)
                if token in p.text:
                    p.text = p.text.replace(token, val)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if token in cell.text:
                        cell.text = cell.text.replace(token, val)


def aplicar_datos_a_docx(doc_path: Path, datos: dict) -> None:
    doc = Document(str(doc_path))
    ghs = datos.get("ghs") or {}
    ident = datos.get("identificacion") or {}
    hoy = datetime.now().strftime("%d-%m-%Y")

    mapping: dict[str, str] = {
        "TITULO": (datos.get("titulo") or ident.get("nombre_comercial") or "PRODUCTO").strip().upper(),
        "REFERENCIA": str(datos.get("referencia") or ident.get("referencia_interna") or ""),
        "FECHA_REVISION": str(datos.get("fecha_revision") or datos.get("fecha_emision") or hoy),
        "FECHA_IMPRESION": str(datos.get("fecha_impresion") or hoy),
        "VERSION": str(datos.get("version") or "1.0/GHS/ES"),
    }

    # Sección 1 con bloque identificación si no hay s01 custom
    if not ghs.get("s01"):
        s01 = (
            "1.1. Identificador del producto\n"
            f"Nombre comercial: {mapping['TITULO']}\n"
            f"Nombre de la sustancia: {ident.get('nombre_inci') or mapping['TITULO']}\n"
            f"Número CAS: {ident.get('cas', '')}\n"
            f"Número CE: {ident.get('numero_ce', '')}\n\n"
            "1.2. Usos pertinentes identificados\n"
            f"{ident.get('usos', 'Materia prima grado cosmético / farmacéutico.')}\n\n"
            "1.3. Datos del proveedor\n"
            "Empresa: MCKENNA GROUP S.A.S.\n"
            "Dirección: Bogotá, Colombia\n"
            "Sitio web: www.mckennagroup.co\n"
            "E-mail: operaciones@mckennagroup.co\n\n"
            "1.4. Teléfono de emergencia\n"
            f"{ident.get('telefono_emergencia', 'Servicio SIT +34 91 562 0420')}\n"
        )
        ghs = dict(ghs)
        ghs["s01"] = s01

    for i in range(1, 17):
        key = f"s{i:02d}"
        if key == "s09":
            continue
        val = ghs.get(key) or ghs.get(str(i)) or ""
        if val:
            mapping[f"S{i:02d}"] = str(val).strip()

    mapping.update(_propiedades_a_mapa(datos))
    _reemplazar_en_doc(doc, mapping)

    # Limpiar placeholders no reemplazados
    for p in doc.paragraphs:
        p.text = re.sub(r"\{\{[A-Z0-9_]+\}\}", "No determinado", p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = re.sub(r"\{\{[A-Z0-9_]+\}\}", "No determinado", cell.text)

    doc.save(str(doc_path))


def folder_id_drive() -> str | None:
    if SDS_DRIVE_FOLDER:
        return SDS_DRIVE_FOLDER
    if not DOCS_DRIVE_PARENT:
        return None
    try:
        from app.services.ficha_tecnica import _drive_service, _resolver_subcarpeta

        service = _drive_service()
        return _resolver_subcarpeta(service, DOCS_DRIVE_PARENT, "SDS")
    except Exception:
        return None


def configuracion_drive() -> dict[str, Any]:
    fid = folder_id_drive()
    email = client_email_servicio()
    instrucciones = []
    if email:
        instrucciones.append(f"1) Comparta la carpeta SDS con {email} como Editor.")
    if TDS_IMPERSONATE:
        instrucciones.append(f"2) Subida delegada como {TDS_IMPERSONATE}")
    else:
        instrucciones.append(
            "2) Configure TDS_DRIVE_IMPERSONATE en .env para subir a Mi unidad compartida."
        )
    return {
        "client_email": email,
        "creds_ok": bool(email),
        "impersonate_email": TDS_IMPERSONATE or None,
        "delegacion_configurada": bool(TDS_IMPERSONATE),
        "parent_folder_id": DOCS_DRIVE_PARENT,
        "folder_id": fid,
        "folder_url": f"https://drive.google.com/drive/folders/{fid}" if fid else None,
        "parent_folder_url": f"https://drive.google.com/drive/folders/{DOCS_DRIVE_PARENT}",
        "instrucciones": " ".join(instrucciones),
        "creds_path": CREDS_PATH,
        "plantilla_referencia_pdf": str(PLANTILLA_REF_PDF) if PLANTILLA_REF_PDF.is_file() else None,
        "formato": "GHS Ventós (SDS ELEMI)",
    }


def generar_desde_datos(
    datos: dict,
    *,
    plantilla: Path | None = None,
    salida: Path | None = None,
    generar_pdf: bool = True,
    subir_drive: bool = False,
    guardar_yaml: str | None = None,
) -> dict[str, Any]:
    if guardar_yaml is not None:
        slug = guardar_yaml or re.sub(
            r"[^a-z0-9_]+",
            "_",
            _normalizar((datos.get("titulo") or "producto")),
        ).strip("_")
        guardar_yaml_datos(datos, slug=f"sds_{slug}" if not slug.startswith("sds_") else slug)

    titulo = (datos.get("titulo") or "PRODUCTO").strip()
    nombre = nombre_archivo_desde_titulo(titulo)
    destino_docx = salida or (SDS_DIR / nombre)

    tpl = plantilla or PLANTILLA_DEFAULT
    if not tpl.exists():
        raise FileNotFoundError(
            f"Plantilla SDS no encontrada: {tpl}. Ejecute: python scripts/crear_plantilla_sds_ghs.py"
        )

    destino_docx.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(tpl, destino_docx)
    aplicar_datos_a_docx(destino_docx, datos)

    resultado: dict[str, Any] = {
        "ok": True,
        "titulo": titulo,
        "docx": str(destino_docx),
        "docx_nombre": destino_docx.name,
    }

    pdf_path: Path | None = None
    if generar_pdf:
        pdf_path = exportar_pdf(destino_docx)
        resultado["pdf"] = str(pdf_path)
        resultado["pdf_nombre"] = pdf_path.name

    if subir_drive:
        uploads = []
        folder = folder_id_drive()
        if folder:
            try:
                uploads.append({"tipo": "docx", **subir_a_drive(destino_docx, folder)})
            except Exception as e:
                uploads.append({"tipo": "docx", "error": str(e)})
            if pdf_path:
                try:
                    uploads.append({"tipo": "pdf", **subir_a_drive(pdf_path, folder)})
                except Exception as e:
                    uploads.append({"tipo": "pdf", "error": str(e)})
        else:
            uploads.append({"tipo": "sds", "error": "Carpeta SDS no configurada en Drive"})
        resultado["drive_uploads"] = uploads

    return resultado


def ruta_descarga_segura(nombre: str) -> Path | None:
    nombre = os.path.basename(nombre or "")
    if not re.match(r"^SDS-.+\.(docx|pdf)$", nombre, re.I):
        return None
    path = (SDS_DIR / nombre).resolve()
    try:
        path.relative_to(SDS_DIR.resolve())
    except ValueError:
        return None
    if path.is_file():
        return path
    return None
