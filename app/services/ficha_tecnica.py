"""
Generación de fichas técnicas McKenna (DOCX + PDF) y subida a Google Drive.
"""

from __future__ import annotations

import io
import json
import os
import re
import shutil
import subprocess
import tempfile
import unicodedata
import zipfile
from pathlib import Path
from typing import Any

import yaml
from docx import Document
from docx.table import Table
from PIL import Image

REPO = Path(__file__).resolve().parents[2]
FICHAS_DIR = REPO / "fichas_word"
DATOS_DIR = FICHAS_DIR / "datos"
PLANTILLAS_DIR = FICHAS_DIR / "plantillas"
CABEZOTES_DIR = FICHAS_DIR / "cabezotes"
DISENO_DIR = REPO / "DISENO CORPORATIVO "
PLANTILLA_DEFAULT = FICHAS_DIR / "FT CAOLIN COLOIDAL.docx"
PLANTILLA_DEFAULT_ID = "default"
CABEZOTE_DEFAULT_ID = "default"
PLANTILLA_YAML = DATOS_DIR / "plantilla_ejemplo.yaml"
_CABEZOTE_EXTS = {".jpg", ".jpeg", ".png", ".webp"}

TDS_PARENT_DEFAULT = os.getenv("TDS_DRIVE_PARENT_ID", "1BTXM8bKCnWVYWTTEmKYxcpaQv1TOoZVs")
TDS_FOLDER_PDF = os.getenv("TDS_DRIVE_FOLDER_PDF", "").strip()
TDS_FOLDER_WORD = os.getenv("TDS_DRIVE_FOLDER_WORD", "").strip()
# Usuario Workspace a impersonar (delegación dominio). Obligatorio para subir a Mi unidad compartida.
TDS_IMPERSONATE = (
    os.getenv("TDS_DRIVE_IMPERSONATE", "").strip()
    or os.getenv("GOOGLE_DRIVE_IMPERSONATE_USER", "").strip()
)
CREDS_PATH = os.getenv(
    "GOOGLE_SERVICE_ACCOUNT_PATH",
    str(REPO / "mi-agente-ubuntu-9043f67d9755.json"),
)

_MIME = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

_folder_cache: dict[str, str] = {}


def client_email_servicio() -> str | None:
    if not os.path.exists(CREDS_PATH):
        return None
    try:
        with open(CREDS_PATH, encoding="utf-8") as f:
            return json.load(f).get("client_email")
    except Exception:
        return None


def _normalizar(s: str) -> str:
    t = unicodedata.normalize("NFD", (s or "").lower())
    t = "".join(c for c in t if unicodedata.category(c) != "Mn")
    return re.sub(r"\s+", " ", t).strip()


def nombre_archivo_desde_titulo(titulo: str) -> str:
    t = unicodedata.normalize("NFD", titulo.upper())
    t = "".join(c for c in t if unicodedata.category(c) != "Mn")
    t = re.sub(r"[^A-Z0-9]+", " ", t).strip()
    return f"FT {' '.join(t.split())}.docx"


def cargar_datos_desde_archivo(path: Path) -> dict:
    raw = path.read_text(encoding="utf-8")
    if path.suffix.lower() == ".json":
        return json.loads(raw)
    return yaml.safe_load(raw) or {}


def plantilla_datos_ejemplo() -> dict:
    if PLANTILLA_YAML.exists():
        return cargar_datos_desde_archivo(PLANTILLA_YAML)
    return {
        "titulo": "NOMBRE DEL PRODUCTO",
        "descripcion": "Descripción del ingrediente…",
        "aplicaciones": ["Uso principal…"],
        "identidad": [["NOMBRE DEL PRODUCTO", "Nombre comercial"]],
        "propiedades": [["Apariencia", "Polvo fino"]],
        "microbiologia": [["E. Coli", "Negativo"]],
        "estabilidad": ["Almacenar en lugar seco…"],
    }


def _slug_seguro(valor: str) -> str:
    return re.sub(r"[^a-zA-Z0-9_-]", "", (valor or "").strip())


def _titulo_desde_archivo(path: Path) -> str:
    stem = path.stem.replace("_", " ").replace("-", " ")
    return stem.strip().title() or path.name


def _asegurar_cabezotes_iniciales() -> None:
    """Crea carpeta cabezotes/ y copia assets por defecto si faltan."""
    CABEZOTES_DIR.mkdir(parents=True, exist_ok=True)
    PLANTILLAS_DIR.mkdir(parents=True, exist_ok=True)

    logo_src = DISENO_DIR / "LOGO MCKENNA.jpg"
    logo_dst = CABEZOTES_DIR / "logo_mckenna.jpg"
    if logo_src.is_file() and not logo_dst.is_file():
        shutil.copy2(logo_src, logo_dst)


def _media_cabezote_en_docx(docx_path: Path) -> str | None:
    """Primera imagen embebida en word/document.xml (cabezote del cuerpo)."""
    if not docx_path.is_file():
        return None
    try:
        with zipfile.ZipFile(docx_path) as z:
            if "word/document.xml" not in z.namelist():
                return None
            doc = z.read("word/document.xml").decode("utf-8", errors="ignore")
            m = re.search(r'r:embed="(rId\d+)"', doc)
            if not m:
                return None
            rid = m.group(1)
            rels_name = "word/_rels/document.xml.rels"
            if rels_name not in z.namelist():
                return None
            rels = z.read(rels_name).decode("utf-8", errors="ignore")
            m2 = re.search(rf'Id="{rid}"[^>]+Target="([^"]+)"', rels)
            if not m2:
                return None
            target = m2.group(1).lstrip("/")
            if target.startswith("media/"):
                return "word/" + target
            if target.startswith("../media/"):
                return "word/" + target.replace("../", "")
            return target if target.startswith("word/") else f"word/{target}"
    except Exception:
        return None


def _imagen_a_jpeg_bytes(path: Path) -> bytes:
    with Image.open(path) as img:
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=92, optimize=True)
        return buf.getvalue()


def _imagen_bytes_rgb(data: bytes) -> tuple[bytes, int, int]:
    """Devuelve (jpeg_rgb_bytes, ancho_px, alto_px). Convierte CMYK→RGB si hace falta."""
    with Image.open(io.BytesIO(data)) as img:
        w, h = img.size
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=92, optimize=True)
        return buf.getvalue(), w, h


def _ajustar_cabezote_en_docx(docx_path: Path) -> None:
    """Convierte la imagen del cabezote a RGB y la reposiciona en la esquina superior derecha."""
    # ── 1. Leer el DOCX ──────────────────────────────────────────────────────
    with zipfile.ZipFile(docx_path) as z:
        names = z.namelist()
        files = {n: z.read(n) for n in names}

    body = files.get("word/document.xml", b"").decode("utf-8", errors="ignore")
    rels_raw = files.get("word/_rels/document.xml.rels", b"").decode("utf-8", errors="ignore")

    # ── 2. Localizar el primer rId de imagen en el body ──────────────────────
    rid_m = re.search(r'r:embed="(rId\d+)"', body)
    if not rid_m:
        return
    rid = rid_m.group(1)

    target_m = re.search(rf'Id="{rid}"[^>]+Target="([^"]+)"', rels_raw)
    if not target_m:
        return
    raw_target = target_m.group(1).lstrip("/")
    if raw_target.startswith("media/"):
        media_key = "word/" + raw_target
    elif raw_target.startswith("../media/"):
        media_key = "word/" + raw_target.replace("../", "")
    else:
        media_key = raw_target if raw_target.startswith("word/") else f"word/{raw_target}"

    if media_key not in files:
        return

    # ── 3. Convertir imagen a RGB ─────────────────────────────────────────────
    jpeg_rgb, img_w, img_h = _imagen_bytes_rgb(files[media_key])
    files[media_key] = jpeg_rgb

    # ── 4. Calcular tamaño EMU (máx 5 cm ancho × 2 cm alto) ──────────────────
    # 1 cm = 360 000 EMU
    max_cx = 1_800_000   # 5 cm
    max_cy =   720_000   # 2 cm
    ratio = img_h / img_w if img_w > 0 else 0.5
    cx = max_cx
    cy = int(cx * ratio)
    if cy > max_cy:
        cy = max_cy
        cx = int(cy / ratio) if ratio > 0 else max_cx

    # ── 5. Parchear el anchor XML ─────────────────────────────────────────────
    # 5a. Posición H → derecha del margen de texto
    body = re.sub(
        r"<wp:positionH[^>]*>.*?</wp:positionH>",
        '<wp:positionH relativeFrom="margin"><wp:align>right</wp:align></wp:positionH>',
        body, count=1, flags=re.DOTALL,
    )
    # 5b. Posición V → borde superior del margen
    body = re.sub(
        r"<wp:positionV[^>]*>.*?</wp:positionV>",
        '<wp:positionV relativeFrom="margin"><wp:posOffset>0</wp:posOffset></wp:positionV>',
        body, count=1, flags=re.DOTALL,
    )
    # 5c. Tamaño wp:extent
    body = re.sub(
        r'<wp:extent\s+cx="[^"]*"\s+cy="[^"]*"/>',
        f'<wp:extent cx="{cx}" cy="{cy}"/>',
        body, count=1,
    )
    # 5d. Tamaño en spPr (a:ext dentro de pic:spPr)
    body = re.sub(
        r'(<a:ext\s+cx=")[^"]+("\s+cy=")[^"]+"',
        rf'\g<1>{cx}\g<2>{cy}"',
        body, count=1,
    )
    # 5e. Sacar de "detrás del documento" para que sea visible
    body = body.replace('behindDoc="1"', 'behindDoc="0"', 1)

    files["word/document.xml"] = body.encode("utf-8")

    # ── 6. Reescribir el DOCX ────────────────────────────────────────────────
    fd, tmp = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    try:
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for name in names:
                zout.writestr(name, files[name])
        shutil.move(tmp, docx_path)
    finally:
        if os.path.exists(tmp):
            try:
                os.unlink(tmp)
            except OSError:
                pass


def _reemplazar_media_docx(docx_path: Path, media_interno: str, imagen: Path) -> None:
    """Sustituye un archivo word/media/* dentro del DOCX."""
    jpeg = _imagen_a_jpeg_bytes(imagen)
    fd, tmp = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = jpeg if item.filename == media_interno else zin.read(item.filename)
                zout.writestr(item, data)
        shutil.move(tmp, docx_path)
    finally:
        if os.path.exists(tmp):
            try:
                os.unlink(tmp)
            except OSError:
                pass


def aplicar_cabezote_a_docx(docx_path: Path, cabezote: Path, *, plantilla_ref: Path | None = None) -> None:
    ref = plantilla_ref or docx_path
    media = _media_cabezote_en_docx(ref)
    if not media:
        raise RuntimeError("La plantilla no tiene imagen de cabezote en el cuerpo del documento")
    _reemplazar_media_docx(docx_path, media, cabezote)


def resolver_plantilla_path(plantilla_id: str | None) -> Path:
    pid = (plantilla_id or PLANTILLA_DEFAULT_ID).strip() or PLANTILLA_DEFAULT_ID
    if pid == PLANTILLA_DEFAULT_ID:
        tpl = PLANTILLA_DEFAULT
    else:
        slug = _slug_seguro(pid)
        if not slug:
            raise ValueError("plantilla_id inválido")
        candidatos = list(PLANTILLAS_DIR.glob(f"{slug}.docx")) + list(PLANTILLAS_DIR.glob(f"{slug}.DOCX"))
        tpl = candidatos[0] if candidatos else PLANTILLAS_DIR / f"{slug}.docx"
    if not tpl.is_file():
        raise FileNotFoundError(f"Plantilla no encontrada: {pid}")
    return tpl


def resolver_cabezote_path(cabezote_id: str | None) -> Path | None:
    cid = (cabezote_id or CABEZOTE_DEFAULT_ID).strip() or CABEZOTE_DEFAULT_ID
    if cid == CABEZOTE_DEFAULT_ID:
        return None
    slug = _slug_seguro(cid)
    if not slug:
        raise ValueError("cabezote_id inválido")
    _asegurar_cabezotes_iniciales()
    for ext in _CABEZOTE_EXTS:
        p = CABEZOTES_DIR / f"{slug}{ext}"
        if p.is_file():
            return p
    raise FileNotFoundError(f"Cabezote no encontrado: {cid}")


def listar_plantillas_docx() -> list[dict[str, str]]:
    _asegurar_cabezotes_iniciales()
    items = [{
        "id": PLANTILLA_DEFAULT_ID,
        "nombre": "McKenna estándar",
        "archivo": PLANTILLA_DEFAULT.name,
        "descripcion": "Formato base (FT Caolín coloidal)",
    }]
    if PLANTILLAS_DIR.is_dir():
        for p in sorted(PLANTILLAS_DIR.glob("*.docx")):
            if p.name.startswith("~$"):
                continue
            upper = p.name.upper()
            if upper.startswith("COA ") or upper.startswith("SDS "):
                continue
            sid = p.stem
            items.append({
                "id": sid,
                "nombre": _titulo_desde_archivo(p),
                "archivo": p.name,
                "descripcion": f"Plantilla personalizada ({p.name})",
            })
    return items


def listar_cabezotes() -> list[dict[str, str]]:
    _asegurar_cabezotes_iniciales()
    items = [{
        "id": CABEZOTE_DEFAULT_ID,
        "nombre": "Cabezote de la plantilla",
        "archivo": "",
        "descripcion": "Usa la imagen incluida en la plantilla Word seleccionada",
    }]
    if CABEZOTES_DIR.is_dir():
        vistos: set[str] = set()
        for ext in sorted(_CABEZOTE_EXTS):
            for p in sorted(CABEZOTES_DIR.glob(f"*{ext}")):
                sid = p.stem
                if sid in vistos:
                    continue
                vistos.add(sid)
                items.append({
                    "id": sid,
                    "nombre": _titulo_desde_archivo(p),
                    "archivo": p.name,
                    "descripcion": p.name,
                })
    return items


def opciones_generacion_ficha() -> dict[str, Any]:
    return {
        "plantillas": listar_plantillas_docx(),
        "cabezotes": listar_cabezotes(),
        "plantillas_dir": str(PLANTILLAS_DIR),
        "cabezotes_dir": str(CABEZOTES_DIR),
        "plantilla_default_id": PLANTILLA_DEFAULT_ID,
        "cabezote_default_id": CABEZOTE_DEFAULT_ID,
    }


def ruta_cabezote_segura(cabezote_id: str) -> Path | None:
    slug = _slug_seguro(cabezote_id)
    if not slug or slug == CABEZOTE_DEFAULT_ID:
        return None
    _asegurar_cabezotes_iniciales()
    for ext in _CABEZOTE_EXTS:
        p = (CABEZOTES_DIR / f"{slug}{ext}").resolve()
        try:
            p.relative_to(CABEZOTES_DIR.resolve())
        except ValueError:
            return None
        if p.is_file():
            return p
    return None


def _slug_cabezote_desde_nombre(nombre: str) -> str:
    slug = re.sub(r"[^a-z0-9_]+", "_", _normalizar(nombre)).strip("_")
    return (slug[:60] or "cabezote")


def guardar_cabezote_subido(file_storage, *, nombre: str | None = None) -> dict[str, str]:
    """Guarda imagen subida en fichas_word/cabezotes/ y devuelve metadatos para el selector."""
    _asegurar_cabezotes_iniciales()
    if file_storage is None or not getattr(file_storage, "filename", None):
        raise ValueError("No se recibió archivo de imagen")

    orig = Path(file_storage.filename)
    ext = orig.suffix.lower()
    if ext not in _CABEZOTE_EXTS:
        permitidos = ", ".join(sorted(_CABEZOTE_EXTS))
        raise ValueError(f"Formato no permitido ({ext or 'sin extensión'}). Use: {permitidos}")

    raw = file_storage.read()
    if not raw:
        raise ValueError("El archivo está vacío")
    if len(raw) > 10 * 1024 * 1024:
        raise ValueError("La imagen supera 10 MB")

    try:
        with Image.open(io.BytesIO(raw)) as probe:
            probe.verify()
        img = Image.open(io.BytesIO(raw))
    except Exception as exc:
        raise ValueError("El archivo no es una imagen válida") from exc

    slug_base = _slug_cabezote_desde_nombre(nombre or orig.stem)
    slug = slug_base
    n = 2
    dest = CABEZOTES_DIR / f"{slug}{ext}"
    while dest.exists():
        slug = f"{slug_base}_{n}"
        dest = CABEZOTES_DIR / f"{slug}{ext}"
        n += 1

    try:
        if ext in (".jpg", ".jpeg"):
            if img.mode not in ("RGB", "L"):
                img = img.convert("RGB")
            img.save(dest, format="JPEG", quality=92, optimize=True)
        elif ext == ".webp":
            img.save(dest, format="WEBP", quality=90, method=4)
        else:
            img.save(dest, format="PNG", optimize=True)
    finally:
        img.close()

    return {
        "id": slug,
        "nombre": _titulo_desde_archivo(dest),
        "archivo": dest.name,
        "descripcion": dest.name,
    }


def listar_yaml_datos() -> list[dict[str, str]]:
    if not DATOS_DIR.is_dir():
        return []
    items = []
    for p in sorted(DATOS_DIR.glob("*.yaml")) + sorted(DATOS_DIR.glob("*.yml")):
        if p.name.startswith("plantilla"):
            continue
        try:
            d = cargar_datos_desde_archivo(p)
            titulo = (d.get("titulo") or p.stem).strip()
        except Exception:
            titulo = p.stem
        items.append({"id": p.stem, "archivo": p.name, "titulo": titulo})
    return items


def _filas_tabla(data: list) -> list[list[str]]:
    filas = []
    for item in data or []:
        if isinstance(item, (list, tuple)) and len(item) >= 2:
            filas.append([str(item[0]).strip(), str(item[1]).strip()])
        elif isinstance(item, dict):
            for k, v in item.items():
                filas.append([str(k).strip(), str(v).strip()])
    return filas


def _rellenar_tabla(tabla: Table, filas: list[list[str]]) -> None:
    for i, (label, valor) in enumerate(filas):
        if i < len(tabla.rows):
            row = tabla.rows[i]
            row.cells[0].text = label
            if len(row.cells) > 1:
                row.cells[-1].text = valor
        else:
            row = tabla.add_row()
            row.cells[0].text = label
            row.cells[-1].text = valor
    for j in range(len(filas), len(tabla.rows)):
        for cell in tabla.rows[j].cells:
            cell.text = ""


def _buscar_parrafo(doc: Document, texto_buscar: str, estilo: str | None = None) -> int | None:
    objetivo = _normalizar(texto_buscar)
    for i, p in enumerate(doc.paragraphs):
        if estilo and p.style.name != estilo:
            continue
        if _normalizar(p.text) == objetivo:
            return i
    return None


def _insertar_despues(doc: Document, idx: int, textos: list[str], estilo: str = "Normal") -> None:
    if idx is None or idx < 0:
        return
    ref = doc.paragraphs[idx]._element
    parent = ref.getparent()
    pos = parent.index(ref) + 1
    for texto in textos:
        nuevo = doc.add_paragraph(texto, style=estilo)
        parent.insert(pos, nuevo._element)
        pos += 1


def _formatear_fecha_revision(fecha: str) -> str:
    """Normaliza a DD MM YYYY (formato habitual en fichas McKenna)."""
    t = (fecha or "").strip()
    if not t:
        return t
    m = re.match(r"^(\d{4})-(\d{2})-(\d{2})$", t)
    if m:
        return f"{m.group(3)} {m.group(2)} {m.group(1)}"
    m = re.match(r"^(\d{2})[-/.](\d{2})[-/.](\d{4})$", t)
    if m:
        return f"{m.group(1)} {m.group(2)} {m.group(3)}"
    return t


def _valor_en_filas(filas: list[list[str]], *claves: str) -> str:
    claves_n = {_normalizar(c) for c in claves}
    for label, val in filas:
        if _normalizar(label) in claves_n and str(val).strip():
            return str(val).strip()
    return ""


def normalizar_datos_ficha(datos: dict) -> dict:
    """Convierte campos estructurados del panel a identidad/propiedades para el DOCX."""
    d = dict(datos or {})
    filas_id_legacy = _filas_tabla(d.get("identidad"))
    filas_prop_legacy = _filas_tabla(d.get("propiedades"))

    nombre = (d.get("nombre_producto") or d.get("titulo") or "").strip()
    if not nombre:
        nombre = _valor_en_filas(filas_id_legacy, "nombre del producto")
    if nombre:
        d["titulo"] = nombre.upper()
        d["nombre_producto"] = nombre

    ref = (d.get("referencia") or "").strip() or _valor_en_filas(
        filas_id_legacy, "referencia siigo", "referencia interna", "referencia"
    )
    sinonimos = (d.get("sinonimos") or "").strip() or _valor_en_filas(filas_id_legacy, "sinonimos", "sinonimo")
    cas = (d.get("cas") or "").strip() or _valor_en_filas(filas_id_legacy, "cas", "cas #")
    fecha = (d.get("fecha_revision") or "").strip() or _valor_en_filas(
        filas_id_legacy, "fecha de revision", "fecha revision"
    )

    identidad_out: list[list[str]] = []
    if nombre:
        identidad_out.append(["NOMBRE DEL PRODUCTO", nombre])
    if ref:
        identidad_out.append(["REFERENCIA SIIGO", ref])
    if sinonimos:
        identidad_out.append(["SINÓNIMOS", sinonimos])
    if cas:
        cas_val = re.sub(r"^(CAS\s*#?\s*)", "", cas, flags=re.I).strip() or cas
        identidad_out.append(["CAS #", cas_val])
    if fecha:
        identidad_out.append(["FECHA DE REVISIÓN", _formatear_fecha_revision(fecha)])

    comp_rows = _filas_tabla(d.get("composicion"))
    comp_labels = {_normalizar(r[0]) for r in comp_rows if r}
    for row in comp_rows:
        if row[0] or row[1]:
            identidad_out.append(row)

    skip_labels = {
        "nombre del producto",
        "referencia siigo",
        "referencia interna",
        "referencia",
        "sinonimos",
        "sinonimo",
        "cas",
        "cas #",
        "fecha de revision",
        "fecha revision",
    }
    skip_labels |= comp_labels
    for row in filas_id_legacy:
        if _normalizar(row[0]) not in skip_labels:
            identidad_out.append(row)

    cf = d.get("caracteristicas_fisicas") if isinstance(d.get("caracteristicas_fisicas"), dict) else {}
    propiedades: list[list[str]] = []
    for key, label in (
        ("apariencia", "Apariencia"),
        ("punto_fusion", "Punto de fusión"),
        ("indice_saponificacion", "Índice de saponificación"),
        ("ph", "pH"),
        ("olor", "Olor"),
        ("sabor", "Sabor"),
        ("formula_quimica", "Fórmula química"),
        ("solubilidad", "Solubilidad"),
    ):
        val = (cf.get(key) or "").strip() if cf else ""
        if not val:
            val = _valor_en_filas(filas_prop_legacy, label.lower(), _normalizar(label))
        if val:
            propiedades.append([label, val])

    extra = d.get("propiedades_extra") or d.get("propiedades_lista") or []
    if isinstance(extra, str):
        extra = [ln.strip() for ln in extra.split("\n") if ln.strip()]
    for item in extra:
        if isinstance(item, str):
            if "|" in item:
                a, b = item.split("|", 1)
                propiedades.append([a.strip(), b.strip()])
            elif item.strip():
                propiedades.append([item.strip(), ""])
        elif isinstance(item, (list, tuple)) and len(item) >= 2:
            propiedades.append([str(item[0]).strip(), str(item[1]).strip()])

    fisicas_labels = {
        "apariencia",
        "punto de fusion",
        "indice de saponificacion",
        "ph",
        "olor",
        "sabor",
        "formula quimica",
        "solubilidad",
    }
    existentes = {_normalizar(p[0]) for p in propiedades}
    for row in filas_prop_legacy:
        ln = _normalizar(row[0])
        if ln not in existentes and ln not in fisicas_labels:
            propiedades.append(row)

    apps = d.get("aplicaciones") or []
    if isinstance(apps, str):
        apps = [ln.strip() for ln in apps.split("\n") if ln.strip()]
    d["referencia"] = ref
    d["sinonimos"] = sinonimos
    d["cas"] = cas
    d["fecha_revision"] = fecha
    d["identidad"] = identidad_out
    d["propiedades"] = propiedades
    d["aplicaciones"] = apps
    return d


def aplicar_datos_a_docx(doc_path: Path, datos: dict) -> None:
    datos = normalizar_datos_ficha(datos)
    doc = Document(str(doc_path))
    titulo = (datos.get("titulo") or "PRODUCTO").strip().upper()

    for p in doc.paragraphs:
        if p.text.strip() and p.style.name == "Normal":
            p.text = titulo
            break

    idx_desc = _buscar_parrafo(doc, "DESCRIPCIÓN", "Heading 1") or _buscar_parrafo(doc, "DESCRIPCION", "Heading 1")
    if idx_desc is not None:
        for j in range(idx_desc + 1, len(doc.paragraphs)):
            p = doc.paragraphs[j]
            if p.style.name == "Heading 1":
                break
            if p.text.strip() and p.style.name == "Normal":
                p.text = (datos.get("descripcion") or "").strip()
                break

    idx_app = _buscar_parrafo(doc, "APLICACIONES")
    idx_prop = _buscar_parrafo(doc, "PROPIEDADES FÍSICO-QUÍMICAS", "Heading 1") or _buscar_parrafo(
        doc, "PROPIEDADES FISICO-QUIMICAS", "Heading 1"
    )
    if idx_app is not None and idx_prop is not None:
        apps = datos.get("aplicaciones") or []
        if isinstance(apps, str):
            apps = [apps]
        for i in range(idx_prop - 1, idx_app, -1):
            p = doc.paragraphs[i]
            if _normalizar(p.text) not in ("aplicaciones",):
                p._element.getparent().remove(p._element)
        _insertar_despues(doc, idx_app, [str(a).strip() for a in apps if str(a).strip()])

    tablas = doc.tables
    filas_id = _filas_tabla(datos.get("identidad"))
    filas_prop = _filas_tabla(datos.get("propiedades"))
    filas_micro = _filas_tabla(datos.get("microbiologia"))

    if len(tablas) >= 1 and filas_id:
        _rellenar_tabla(tablas[0], filas_id)
    if len(tablas) >= 2 and filas_prop:
        _rellenar_tabla(tablas[1], filas_prop)
    if len(tablas) >= 3 and filas_micro:
        _rellenar_tabla(tablas[2], filas_micro)

    nota = (datos.get("nota_micro") or "").strip()
    idx_micro_h = _buscar_parrafo(doc, "PROPIEDADES MICROBIOLÓGICAS") or _buscar_parrafo(
        doc, "PROPIEDADES MICROBIOLOGICAS"
    )
    if idx_micro_h is not None and nota:
        for j in range(idx_micro_h + 1, len(doc.paragraphs)):
            p = doc.paragraphs[j]
            if p.style.name == "Heading 1":
                break
            if p.style.name == "Normal" and "nota" in _normalizar(p.text):
                p.text = f"Nota: {nota}"
                break
        else:
            _insertar_despues(doc, idx_micro_h, [f"Nota: {nota}"])

    idx_est = _buscar_parrafo(doc, "ESTABILIDAD Y ALMACENAMIENTO", "Heading 1")
    estab = datos.get("estabilidad") or []
    if isinstance(estab, str):
        estab = [estab]
    if idx_est is not None and estab:
        for j in range(idx_est + 1, len(doc.paragraphs)):
            p = doc.paragraphs[j]
            if p.style.name == "Heading 1":
                break
            if p.text.strip() and p.style.name == "Normal":
                p.text = str(estab[0]).strip()
                for extra in estab[1:]:
                    _insertar_despues(doc, j, [str(extra).strip()])
                break

    doc.save(str(doc_path))


def exportar_pdf(docx_path: Path, pdf_dir: Path | None = None) -> Path:
    out_dir = pdf_dir or docx_path.parent
    cmd = [
        "libreoffice",
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(out_dir),
        str(docx_path),
    ]
    proc = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
    if proc.returncode != 0:
        err = (proc.stderr or proc.stdout or "").strip() or f"código {proc.returncode}"
        raise RuntimeError(f"LibreOffice no generó PDF: {err}")
    pdf_path = out_dir / (docx_path.stem + ".pdf")
    if not pdf_path.exists():
        raise FileNotFoundError(f"No se generó PDF en {out_dir}")
    return pdf_path


def _drive_scopes() -> list[str]:
    if TDS_IMPERSONATE:
        return ["https://www.googleapis.com/auth/drive"]
    return ["https://www.googleapis.com/auth/drive.file"]


def _drive_service():
    from google.oauth2 import service_account
    from googleapiclient.discovery import build

    if not os.path.exists(CREDS_PATH):
        raise FileNotFoundError(f"Credenciales no encontradas: {CREDS_PATH}")
    creds = service_account.Credentials.from_service_account_file(
        CREDS_PATH,
        scopes=_drive_scopes(),
    )
    if TDS_IMPERSONATE:
        creds = creds.with_subject(TDS_IMPERSONATE)
    return build("drive", "v3", credentials=creds, cache_discovery=False)


def _mensaje_error_drive(exc: Exception) -> str:
    txt = str(exc)
    if "storageQuotaExceeded" in txt or "do not have storage quota" in txt.lower():
        return (
            "La cuenta de servicio no puede guardar en Mi unidad (sin cuota). "
            "Configure TDS_DRIVE_IMPERSONATE con un correo @mckennagroup.co del dueño "
            "de las carpetas y active delegación de dominio en Google Workspace "
            "(ver panel Fichas técnicas → ayuda Drive)."
        )
    if "insufficientParentPermissions" in txt:
        return (
            f"Sin permiso en la carpeta. Comparta WORD y PDF con "
            f"{client_email_servicio()} como Editor."
        )
    return txt


def _resolver_subcarpeta(service, parent_id: str, nombre: str) -> str | None:
    clave = f"{parent_id}:{nombre.upper()}"
    if clave in _folder_cache:
        return _folder_cache[clave]
    q = (
        f"'{parent_id}' in parents and mimeType = 'application/vnd.google-apps.folder' "
        f"and trashed = false"
    )
    res = service.files().list(q=q, fields="files(id, name)", pageSize=100).execute()
    objetivo = nombre.strip().upper()
    for f in res.get("files", []):
        if (f.get("name") or "").strip().upper() == objetivo:
            _folder_cache[clave] = f["id"]
            return f["id"]
    return None


def folder_id_por_tipo(tipo: str) -> str | None:
    """tipo: 'pdf' | 'word' | 'docx'"""
    t = tipo.lower()
    if t in ("pdf",):
        explicit = TDS_FOLDER_PDF
        subnombre = "PDF"
    else:
        explicit = TDS_FOLDER_WORD
        subnombre = "WORD"
    if explicit:
        return explicit
    if not TDS_PARENT_DEFAULT:
        return None
    try:
        service = _drive_service()
        return _resolver_subcarpeta(service, TDS_PARENT_DEFAULT, subnombre)
    except Exception:
        return None


def configuracion_drive() -> dict[str, Any]:
    email = client_email_servicio()
    pdf_id = folder_id_por_tipo("pdf")
    word_id = folder_id_por_tipo("word")
    delegacion_ok = bool(TDS_IMPERSONATE)
    instrucciones = []
    if email:
        instrucciones.append(
            f"1) Comparta carpetas WORD y PDF con {email} como Editor."
        )
    if not delegacion_ok:
        instrucciones.append(
            "2) En .env defina TDS_DRIVE_IMPERSONATE=correo@mckennagroup.co "
            "(dueño de las carpetas) y active delegación de dominio en Admin Google."
        )
    else:
        instrucciones.append(f"2) Subida como usuario delegado: {TDS_IMPERSONATE}")
    return {
        "client_email": email,
        "creds_ok": bool(email),
        "impersonate_email": TDS_IMPERSONATE or None,
        "delegacion_configurada": delegacion_ok,
        "parent_folder_id": TDS_PARENT_DEFAULT,
        "folder_pdf_id": pdf_id,
        "folder_word_id": word_id,
        "folder_pdf_url": f"https://drive.google.com/drive/folders/{pdf_id}" if pdf_id else None,
        "folder_word_url": f"https://drive.google.com/drive/folders/{word_id}" if word_id else None,
        "parent_folder_url": f"https://drive.google.com/drive/folders/{TDS_PARENT_DEFAULT}",
        "instrucciones": " ".join(instrucciones) if instrucciones else "Configure credenciales Google.",
        "ayuda_delegacion": (
            "Google Cloud → cuenta de servicio → Delegación en todo el dominio. "
            "Admin Workspace → Controles de API → Delegación: añadir Client ID de la SA "
            "con alcance https://www.googleapis.com/auth/drive"
        ),
    }


def subir_a_drive(archivo: Path, folder_id: str) -> dict[str, str]:
    from googleapiclient.http import MediaFileUpload

    if not folder_id:
        raise ValueError("Carpeta Drive no configurada o sin permisos (WORD/PDF)")
    service = _drive_service()
    ext = archivo.suffix.lower()
    mime = _MIME.get(ext, "application/octet-stream")
    meta = {"name": archivo.name, "parents": [folder_id]}
    media = MediaFileUpload(str(archivo), mimetype=mime, resumable=True)
    try:
        created = (
            service.files()
            .create(
                body=meta,
                media_body=media,
                fields="id, webViewLink",
                supportsAllDrives=True,
            )
            .execute()
        )
    except Exception as exc:
        raise RuntimeError(_mensaje_error_drive(exc)) from exc
    fid = created.get("id", "")
    return {
        "file_id": fid,
        "webViewLink": created.get("webViewLink") or f"https://drive.google.com/file/d/{fid}/view",
        "nombre": archivo.name,
    }


def guardar_yaml_datos(datos: dict, slug: str | None = None) -> Path:
    titulo = (datos.get("titulo") or "producto").strip()
    if slug:
        nombre = re.sub(r"[^a-z0-9_]+", "_", _normalizar(slug))
    else:
        nombre = re.sub(r"[^a-z0-9_]+", "_", _normalizar(titulo))
    nombre = nombre.strip("_") or "producto"
    DATOS_DIR.mkdir(parents=True, exist_ok=True)
    path = DATOS_DIR / f"{nombre}.yaml"
    path.write_text(
        yaml.dump(datos, allow_unicode=True, sort_keys=False, default_flow_style=False),
        encoding="utf-8",
    )
    return path


FICHAS_PDF_DIR = FICHAS_DIR / "pdf"


def _cabezote_src_html(cabezote_id: str | None) -> str | None:
    """Devuelve data URL base64 del cabezote para el template HTML.
    Sirve los bytes originales sin re-encodear para preservar calidad;
    solo convierte si el modo es CMYK o similar."""
    import base64

    path = resolver_cabezote_path(cabezote_id)
    if not path:
        for nombre in ("logo_mckenna.jpg", "mckenna_estandar.jpg", "isotipo_mckenna.png"):
            candidato = CABEZOTES_DIR / nombre
            if candidato.is_file():
                path = candidato
                break
    if not path or not path.is_file():
        return None
    try:
        raw = path.read_bytes()
        ext = path.suffix.lower()

        # Detectar si necesita conversión de modo (CMYK → RGB)
        with Image.open(io.BytesIO(raw)) as probe:
            modo = probe.mode

        if modo in ("RGB", "L", "RGBA", "P"):
            # Servir bytes originales sin pérdida adicional
            mime = "image/png" if ext == ".png" else "image/jpeg"
            b64 = base64.b64encode(raw).decode()
        else:
            # CMYK u otro modo: convertir a PNG lossless
            with Image.open(io.BytesIO(raw)) as img:
                img = img.convert("RGBA" if ext == ".png" else "RGB")
                buf = io.BytesIO()
                fmt = "PNG" if ext == ".png" else "JPEG"
                kw = {} if fmt == "PNG" else {"quality": 98, "subsampling": 0}
                img.save(buf, format=fmt, **kw)
            mime = "image/png" if ext == ".png" else "image/jpeg"
            b64 = base64.b64encode(buf.getvalue()).decode()

        return f"data:{mime};base64,{b64}"
    except Exception:
        return None


def _contexto_html(datos: dict, cabezote_id: str | None = None) -> dict:
    """Transforma datos normalizados en variables para el template Jinja2."""
    d = normalizar_datos_ficha(datos)

    titulo = (d.get("titulo") or d.get("nombre_producto") or "PRODUCTO").strip().upper()
    referencia = (d.get("referencia") or "").strip()
    sinonimos = (d.get("sinonimos") or "").strip()
    cas = (d.get("cas") or "").strip()
    fecha = _formatear_fecha_revision(d.get("fecha_revision") or "")
    from app.services.documento_cientifico import _asegurar_punto_final
    descripcion = _asegurar_punto_final((d.get("descripcion") or "").strip())

    # Características físico-químicas (campos fijos del formulario)
    fisicas_keys = {
        "apariencia", "punto de fusion", "indice de saponificacion", "ph", "olor", "sabor",
        "formula quimica", "solubilidad",
    }
    cf = d.get("caracteristicas_fisicas") or {}
    propiedades_fijas: list[tuple[str, str]] = []
    etiquetas = [
        ("apariencia", "Apariencia"),
        ("punto_fusion", "Punto de fusión"),
        ("indice_saponificacion", "Índice de saponificación"),
        ("ph", "pH"),
        ("olor", "Olor"),
        ("sabor", "Sabor"),
        ("formula_quimica", "Fórmula química"),
        ("solubilidad", "Solubilidad"),
    ]
    for key, label in etiquetas:
        val = (cf.get(key) or "").strip()
        if not val:
            val = _valor_en_filas(_filas_tabla(d.get("propiedades")), label.lower(), _normalizar(label))
        if val:
            propiedades_fijas.append((label, val))

    # Propiedades extra (del campo propiedades_lista)
    propiedades_extra: list[tuple[str, str]] = []
    extra_raw = d.get("propiedades") or []
    for label, val in _filas_tabla(extra_raw):
        ln = _normalizar(label)
        if ln not in fisicas_keys and val:
            propiedades_extra.append((label, val))

    # Aplicaciones
    apps = d.get("aplicaciones") or []
    if isinstance(apps, str):
        apps = [a.strip() for a in apps.split("\n") if a.strip()]
    apps = [_asegurar_punto_final(a) for a in apps]

    # Composición
    composicion = [(r[0], r[1]) for r in _filas_tabla(d.get("composicion")) if r[0] or r[1]]

    # Modo de uso (sección propia después de Aplicaciones)
    modo_uso = _asegurar_punto_final((d.get("modo_uso") or "").strip())

    # Recomendaciones GHS
    recomendaciones_raw = (d.get("recomendaciones") or "").strip()
    recomendaciones = [
        _asegurar_punto_final(l.strip()) for l in recomendaciones_raw.split("\n") if l.strip()
    ]

    # Lote
    lote = (d.get("lote") or "").strip()

    # Color de acento (tema)
    _hex_re = re.compile(r'^#[0-9A-Fa-f]{6}$')
    color_raw = (d.get("color_acento") or "").strip()
    color_acento = color_raw if _hex_re.match(color_raw) else "#069DC2"

    return {
        "titulo": titulo,
        "referencia": referencia,
        "sinonimos": sinonimos,
        "cas": cas,
        "fecha_revision": fecha,
        "descripcion": descripcion,
        "propiedades": propiedades_fijas,
        "propiedades_extra": propiedades_extra,
        "aplicaciones": apps,
        "modo_uso": modo_uso,
        "recomendaciones": recomendaciones,
        "lote": lote,
        "color_acento": color_acento,
        "composicion": composicion,
        "cabezote_src": _cabezote_src_html(cabezote_id),
    }


def generar_pdf_html(
    datos: dict,
    *,
    cabezote_id: str | None = None,
    salida: Path | None = None,
) -> dict:
    """Genera PDF directo desde datos del formulario usando WeasyPrint (sin Word)."""
    from jinja2 import Environment, FileSystemLoader
    from weasyprint import HTML, CSS

    FICHAS_PDF_DIR.mkdir(parents=True, exist_ok=True)

    ctx = _contexto_html(datos, cabezote_id)
    titulo = ctx["titulo"]
    nombre_pdf = nombre_archivo_desde_titulo(titulo).replace(".docx", ".pdf")
    destino = salida or (FICHAS_PDF_DIR / nombre_pdf)

    tpl_dir = Path(__file__).resolve().parents[1] / "templates"
    env = Environment(loader=FileSystemLoader(str(tpl_dir)), autoescape=True)
    tpl = env.get_template("ficha_tecnica_pdf.html")
    html_str = tpl.render(**ctx)

    HTML(string=html_str, base_url=str(tpl_dir)).write_pdf(str(destino))

    return {
        "ok": True,
        "titulo": titulo,
        "pdf": str(destino),
        "pdf_nombre": destino.name,
        "docx": "",
        "docx_nombre": "",
    }


COMPLETO_PDF_DIR = FICHAS_DIR / "completo"


def _contexto_coa(datos_coa: dict) -> dict:
    """Aplana los datos del formulario COA para el template HTML combinado."""
    ident = (datos_coa.get("identificacion") or {})
    lote = (datos_coa.get("lote") or {})
    emp = (datos_coa.get("empaque") or {})
    parametros_raw = datos_coa.get("parametros") or []
    filas: list[list[str]] = []
    if isinstance(parametros_raw, list):
        for item in parametros_raw:
            if isinstance(item, dict):
                filas.append([str(item.get("parametro", "") or item.get(list(item.keys())[0], "")),
                               str(item.get("especificacion", "") or (list(item.values())[1] if len(item) > 1 else "")),
                               str(item.get("resultado", "") or (list(item.values())[2] if len(item) > 2 else ""))])
            elif isinstance(item, (list, tuple)):
                filas.append([str(c) for c in item] + [""] * max(0, 3 - len(item)))
            elif isinstance(item, str):
                partes = item.split("|")
                filas.append([str(p.strip()) for p in partes] + [""] * max(0, 3 - len(partes)))
    return {
        "titulo": (datos_coa.get("titulo") or "").strip().upper() or None,
        "referencia": (ident.get("referencia_interna") or datos_coa.get("referencia") or "").strip(),
        "nombre_comercial": (ident.get("nombre_comercial") or "").strip(),
        "inci": (ident.get("nombre_inci") or "").strip(),
        "cas": (ident.get("cas") or "").strip(),
        "formula": (ident.get("formula_molecular") or "").strip(),
        "einces": (ident.get("einces") or "").strip(),
        "concentracion": (ident.get("concentracion") or "").strip(),
        "grado": (ident.get("grado") or "").strip(),
        "presentacion": (ident.get("presentacion") or "").strip(),
        "incluye": (ident.get("incluye") or "").strip(),
        "lote_numero": (lote.get("numero") or "").strip(),
        "lote_fab": (lote.get("fecha_fabricacion") or "").strip(),
        "lote_venc": (lote.get("fecha_vencimiento") or "").strip(),
        "vida_util": (lote.get("vida_util") or "").strip(),
        "tamano_lote": (lote.get("tamano_lote") or "").strip(),
        "pais_origen": (lote.get("pais_origen") or "").strip(),
        "fecha_analisis": (lote.get("fecha_analisis") or "").strip(),
        "fecha_emision": (lote.get("fecha_emision") or "").strip(),
        "parametros": filas,
        "empaque": (emp.get("empaque_original") or "").strip(),
        "almacenamiento": (emp.get("almacenamiento") or "").strip(),
        "precauciones": (emp.get("precauciones") or "").strip(),
        "observaciones": (emp.get("observaciones") or "").strip(),
        "codigo_verificacion": (datos_coa.get("codigo_verificacion") or "").strip(),
    }


_COA_CAMPOS_EXCLUSIVOS = (
    "einces", "concentracion", "grado", "presentacion", "incluye",
    "lote_numero", "lote_fab", "lote_venc", "vida_util", "tamano_lote",
    "pais_origen", "fecha_analisis", "fecha_emision",
    "empaque", "almacenamiento", "precauciones", "observaciones",
    "codigo_verificacion",
)


def _coa_diligenciado(coa_ctx: dict) -> bool:
    """True si el COA trae contenido propio (más allá de lo que ya mirror la FT: nombre, INCI, CAS…)."""
    if any((coa_ctx.get(campo) or "").strip() for campo in _COA_CAMPOS_EXCLUSIVOS):
        return True
    for fila in coa_ctx.get("parametros") or []:
        if any((celda or "").strip() for celda in fila):
            return True
    return False


def _contexto_sds(datos_sds: dict) -> dict:
    """Aplana los datos del formulario SDS para el template HTML combinado."""
    ident = (datos_sds.get("identificacion") or {})
    pel = (datos_sds.get("peligros") or {})
    man = (datos_sds.get("manipulacion") or {})
    reg = (datos_sds.get("regulatorio") or {})

    def _filas3(raw) -> list[list[str]]:
        filas: list[list[str]] = []
        if not raw:
            return filas
        if isinstance(raw, list):
            for item in raw:
                if isinstance(item, dict):
                    vals = [str(v) for v in item.values()]
                    filas.append((vals + ["", "", ""])[:3])
                elif isinstance(item, (list, tuple)):
                    filas.append([str(c) for c in item] + [""] * max(0, 3 - len(item)))
                elif isinstance(item, str):
                    partes = item.split("|")
                    filas.append([p.strip() for p in partes] + [""] * max(0, 3 - len(partes)))
        elif isinstance(raw, str):
            for linea in raw.split("\n"):
                linea = linea.strip()
                if linea:
                    partes = linea.split("|")
                    filas.append([p.strip() for p in partes] + [""] * max(0, 3 - len(partes)))
        return filas

    def _filas2(raw) -> list[tuple[str, str]]:
        filas: list[tuple[str, str]] = []
        if not raw:
            return filas
        if isinstance(raw, list):
            for item in raw:
                if isinstance(item, dict):
                    vals = list(item.values())
                    filas.append((str(vals[0]) if vals else "", str(vals[1]) if len(vals) > 1 else ""))
                elif isinstance(item, (list, tuple)):
                    filas.append((str(item[0]) if item else "", str(item[1]) if len(item) > 1 else ""))
                elif isinstance(item, str):
                    partes = item.split("|", 1)
                    filas.append((partes[0].strip(), partes[1].strip() if len(partes) > 1 else ""))
        elif isinstance(raw, str):
            for linea in raw.split("\n"):
                linea = linea.strip()
                if linea:
                    partes = linea.split("|", 1)
                    filas.append((partes[0].strip(), partes[1].strip() if len(partes) > 1 else ""))
        return filas

    return {
        "titulo": (datos_sds.get("titulo") or "").strip().upper() or None,
        "referencia": (ident.get("referencia_interna") or datos_sds.get("referencia") or "").strip(),
        "nombre_comercial": (ident.get("nombre_comercial") or "").strip(),
        "inci": (ident.get("nombre_inci") or "").strip(),
        "cas": (ident.get("cas") or "").strip(),
        "formula": (ident.get("formula_molecular") or "").strip(),
        "usos": (ident.get("usos") or "").strip(),
        "telefono": (ident.get("telefono_emergencia") or "").strip(),
        "clasificacion": (pel.get("clasificacion") or "").strip(),
        "pictogramas": (pel.get("pictogramas") or "").strip(),
        "composicion": _filas3(datos_sds.get("composicion")),
        "primeros_auxilios": _filas2(datos_sds.get("primeros_auxilios")),
        "manipulacion": (man.get("manipulacion") or "").strip(),
        "almacenamiento": (man.get("almacenamiento") or "").strip(),
        "propiedades": _filas2(datos_sds.get("propiedades")),
        "normativa": (reg.get("normativa") or "").strip(),
        "observaciones": (reg.get("observaciones") or "").strip(),
    }


_SDS_CAMPOS_EXCLUSIVOS = (
    "usos", "telefono", "clasificacion", "pictogramas",
    "manipulacion", "almacenamiento", "normativa", "observaciones",
)


def _sds_diligenciado(sds_ctx: dict) -> bool:
    """True si el SDS trae contenido propio (más allá de lo que ya mirror la FT: nombre, INCI, CAS…)."""
    if any((sds_ctx.get(campo) or "").strip() for campo in _SDS_CAMPOS_EXCLUSIVOS):
        return True
    for clave in ("composicion", "primeros_auxilios", "propiedades"):
        for fila in sds_ctx.get(clave) or []:
            if any((celda or "").strip() for celda in fila):
                return True
    return False


def generar_pdf_completo(
    datos_ft: dict,
    datos_coa: dict | None = None,
    datos_sds: dict | None = None,
    *,
    cabezote_id: str | None = None,
    salida: Path | None = None,
) -> dict:
    """Genera un PDF unificado FT + COA + SDS desde datos de formulario."""
    from jinja2 import Environment, FileSystemLoader
    from weasyprint import HTML

    COMPLETO_PDF_DIR.mkdir(parents=True, exist_ok=True)

    ft_ctx = _contexto_html(datos_ft, cabezote_id)
    titulo = ft_ctx["titulo"]

    nombre_pdf = nombre_archivo_desde_titulo(titulo).replace(".docx", ".pdf")
    # "FT INULINA.pdf" → quitar prefijo "FT " y agregar "FT COA SDS "
    nombre_pdf = re.sub(r"^FT\s+", "", nombre_pdf, flags=re.I)
    nombre_pdf = f"FT COA SDS {nombre_pdf}"
    destino = salida or (COMPLETO_PDF_DIR / nombre_pdf)

    coa_ctx = _contexto_coa(datos_coa) if datos_coa else None
    if coa_ctx and not _coa_diligenciado(coa_ctx):
        coa_ctx = None
    sds_ctx = _contexto_sds(datos_sds) if datos_sds else None
    if sds_ctx and not _sds_diligenciado(sds_ctx):
        sds_ctx = None

    import re as _re
    def _formula_sub(val: str) -> str:
        return _re.sub(r"(\d+)", r"<sub>\1</sub>", str(val or ""))

    tpl_dir = Path(__file__).resolve().parents[1] / "templates"
    env = Environment(loader=FileSystemLoader(str(tpl_dir)), autoescape=True)
    env.filters["formula_sub"] = _formula_sub
    tpl = env.get_template("documento_completo_pdf.html")
    html_str = tpl.render(
        titulo=titulo,
        color_acento=ft_ctx["color_acento"],
        cabezote_src=ft_ctx["cabezote_src"],
        ft=ft_ctx,
        coa=coa_ctx,
        sds=sds_ctx,
    )

    HTML(string=html_str, base_url=str(tpl_dir)).write_pdf(str(destino))

    return {
        "ok": True,
        "titulo": titulo,
        "pdf": str(destino),
        "pdf_nombre": destino.name,
        "docx": "",
        "docx_nombre": "",
    }


def generar_desde_datos(
    datos: dict,
    *,
    plantilla: Path | None = None,
    plantilla_id: str | None = None,
    cabezote_id: str | None = None,
    salida: Path | None = None,
    generar_pdf: bool = True,
    subir_drive: bool = False,
    guardar_yaml: str | None = None,
) -> dict[str, Any]:
    if guardar_yaml is not None:
        guardar_yaml_datos(normalizar_datos_ficha(datos), slug=guardar_yaml or None)

    titulo = (datos.get("titulo") or "PRODUCTO").strip()
    nombre = nombre_archivo_desde_titulo(titulo)
    destino_docx = salida or (FICHAS_DIR / nombre)

    if plantilla is not None:
        tpl = plantilla
    else:
        tpl = resolver_plantilla_path(plantilla_id)
    if not tpl.exists():
        raise FileNotFoundError(f"Plantilla no encontrada: {tpl}")

    destino_docx.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(tpl, destino_docx)

    cabezote_path = resolver_cabezote_path(cabezote_id)
    if cabezote_path:
        aplicar_cabezote_a_docx(destino_docx, cabezote_path, plantilla_ref=tpl)

    # Convierte cabezote a RGB (elimina CMYK) y lo reposiciona top-right
    _ajustar_cabezote_en_docx(destino_docx)

    datos_norm = normalizar_datos_ficha(datos)
    aplicar_datos_a_docx(destino_docx, datos_norm)

    resultado: dict[str, Any] = {
        "ok": True,
        "titulo": titulo,
        "docx": str(destino_docx),
        "docx_nombre": destino_docx.name,
        "plantilla_id": plantilla_id or PLANTILLA_DEFAULT_ID,
        "cabezote_id": cabezote_id or CABEZOTE_DEFAULT_ID,
    }

    pdf_path: Path | None = None
    if generar_pdf:
        pdf_path = exportar_pdf(destino_docx)
        resultado["pdf"] = str(pdf_path)
        resultado["pdf_nombre"] = pdf_path.name

    if subir_drive:
        uploads = []
        word_folder = folder_id_por_tipo("word")
        if word_folder:
            try:
                uploads.append({"tipo": "word", **subir_a_drive(destino_docx, word_folder)})
            except Exception as e:
                uploads.append({"tipo": "word", "error": str(e)})
        else:
            uploads.append({"tipo": "word", "error": "Carpeta WORD no encontrada en Drive"})

        if pdf_path:
            pdf_folder = folder_id_por_tipo("pdf")
            if pdf_folder:
                try:
                    uploads.append({"tipo": "pdf", **subir_a_drive(pdf_path, pdf_folder)})
                except Exception as e:
                    uploads.append({"tipo": "pdf", "error": str(e)})
            else:
                uploads.append({"tipo": "pdf", "error": "Carpeta PDF no encontrada en Drive"})
        resultado["drive_uploads"] = uploads

    return resultado


def generar_desde_archivo(
    datos_path: Path,
    **kwargs: Any,
) -> dict[str, Any]:
    datos = cargar_datos_desde_archivo(datos_path)
    return generar_desde_datos(datos, **kwargs)


def ruta_descarga_segura(nombre: str) -> Path | None:
    """Archivos FT *.pdf bajo fichas_word/ o fichas_word/pdf/."""
    nombre = os.path.basename(nombre or "")
    if not re.match(r"^FT .+\.pdf$", nombre, re.I):
        return None
    for directorio in (FICHAS_PDF_DIR, FICHAS_DIR):
        path = (directorio / nombre).resolve()
        try:
            path.relative_to(directorio.resolve())
        except ValueError:
            continue
        if path.is_file():
            return path
    return None


_PREFIJO_COMPLETO = re.compile(r"^FT COA SDS[\s\-].+\.(pdf|docx)$", re.I)
_PREFIJO_FT_SIMPLE = re.compile(r"^FT (?!COA\s+SDS).+\.(pdf|docx)$", re.I)
# Compat: nombre histórico usado por rutas que filtraban solo completos
_PREFIJOS_BIBLIOTECA = _PREFIJO_COMPLETO


def _es_nombre_biblioteca(nombre: str) -> bool:
    """True si el nombre es FT simple o documento completo FT+COA+SDS."""
    return bool(_PREFIJO_COMPLETO.match(nombre) or _PREFIJO_FT_SIMPLE.match(nombre))


def listar_archivos_generados() -> list[dict]:
    """Lista PDFs/DOCX de la biblioteca: completos (completo/) y FT simples (pdf/).

    No incluye los Word fuente históricos de fichas_word/ (solo PDFs generados).
    """
    resultado: list[dict] = []
    vistos: set[str] = set()

    # (directorio, patrón, categoría, extensiones permitidas)
    fuentes: list[tuple[Path, re.Pattern[str], str, set[str]]] = [
        (COMPLETO_PDF_DIR, _PREFIJO_COMPLETO, "completo", {".pdf", ".docx"}),
        (FICHAS_PDF_DIR, _PREFIJO_FT_SIMPLE, "ft", {".pdf"}),
        # Legado: algunos FT PDF quedaron en la raíz (nunca listar .docx fuente)
        (FICHAS_DIR, _PREFIJO_FT_SIMPLE, "ft", {".pdf"}),
    ]

    for directorio, patron, categoria, exts in fuentes:
        if not directorio.exists():
            continue
        for p in sorted(directorio.iterdir()):
            if not p.is_file() or p.name.startswith("~") or p.name in vistos:
                continue
            if p.suffix.lower() not in exts:
                continue
            if not patron.match(p.name):
                continue
            # No listar de nuevo un FT de la raíz si ya está en pdf/
            if directorio == FICHAS_DIR and (FICHAS_PDF_DIR / p.name).is_file():
                continue
            vistos.add(p.name)
            stat = p.stat()
            resultado.append({
                "nombre": p.name,
                "tipo": p.suffix.lstrip(".").lower(),
                "tamano": stat.st_size,
                "fecha": int(stat.st_mtime),
                "categoria": categoria,
            })
    return sorted(resultado, key=lambda x: x["nombre"].lower())


def extraer_datos_desde_pdf_ft(path: Path) -> dict:
    """Extrae los campos de una FT generada con WeasyPrint usando PyMuPDF como fallback."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return {}
    try:
        doc = fitz.open(str(path))
        all_text = "\n".join(page.get_text() for page in doc)
    except Exception:
        return {}

    _footer_re = re.compile(r"^\d{4}\s*[·•]\s*P[aá]g\.", re.I)
    lines = [l.strip() for l in all_text.split("\n")]
    lines = [
        l for l in lines
        if l and l != "•" and not l.startswith("McKenna Group S.A.S.") and not _footer_re.match(l)
    ]

    def _n(s: str) -> str:
        return unicodedata.normalize("NFD", s.lower()).encode("ascii", "ignore").decode().strip()

    SEC_KEYS: dict[str, list[str]] = {
        "descripcion":    ["descripcion", "descripción"],
        "caracteristicas": ["caracteristicas fisico-quimicas", "caracteristicas fisico quimicas"],
        "beneficios":     ["beneficios"],
        "aplicaciones":   ["aplicaciones"],
        "modo_uso":       ["modo de uso"],
        "composicion":    ["composicion", "composicion"],
        "ghs":            ["recomendaciones para manejo seguro", "recomendaciones ghs"],
    }
    sec_pos: dict[str, int] = {}
    for i, line in enumerate(lines):
        ln = _n(line)
        for sec, keywords in SEC_KEYS.items():
            if sec not in sec_pos and any(kw in ln for kw in keywords):
                sec_pos[sec] = i

    def _sec_lines(name: str) -> list[str]:
        if name not in sec_pos:
            return []
        start = sec_pos[name] + 1
        nexts = [v for k, v in sec_pos.items() if v > sec_pos[name]]
        end = min(nexts) if nexts else len(lines)
        return [l for l in lines[start:end] if l]

    # --- Encabezado ---
    first = min(sec_pos.values()) if sec_pos else len(lines)
    hdr = lines[:first]
    titulo = ""
    sinonimos_parts: list[str] = []
    cas = ""
    fecha_revision = ""
    referencia = ""
    for line in hdr:
        ln = _n(line)
        if ln.startswith("ficha t") or ln == "ficha tecnica":
            continue
        if ln.startswith("ref."):
            referencia = line[4:].strip()
            continue
        if line.startswith("CAS:"):
            cas = line[4:].strip()
            continue
        if line.startswith("Revisión:"):
            fecha_revision = line[9:].strip()
            continue
        if not titulo:
            titulo = line
        else:
            sinonimos_parts.append(line)
    sinonimos = " ".join(sinonimos_parts)

    # "21 06 2026" -> "2026-06-21"
    m_f = re.match(r'^(\d{1,2})\s+(\d{1,2})\s+(\d{4})$', fecha_revision.strip())
    if m_f:
        fecha_revision = f"{m_f.group(3)}-{m_f.group(2).zfill(2)}-{m_f.group(1).zfill(2)}"

    # --- Descripción ---
    descripcion = " ".join(_sec_lines("descripcion"))

    # --- Características físico-químicas ---
    PROP_MAP: dict[str, str] = {
        "apariencia": "apariencia",
        "punto de fusion": "punto_fusion",
        "punto de fusión": "punto_fusion",
        "indice de saponificacion": "indice_saponificacion",
        "índice de saponificación": "indice_saponificacion",
        "ph": "ph",
        "olor": "olor",
        "formula quimica": "formula_quimica",
        "fórmula química": "formula_quimica",
        "solubilidad": "solubilidad",
        "humedad": "humedad",
        "inercia quimica": "inercia_quimica",
        "inercia química": "inercia_quimica",
    }
    cf: dict[str, str] = {}
    cur_key: str | None = None
    cur_val: list[str] = []

    def _flush_cf() -> None:
        if cur_key and cur_val:
            cf[cur_key] = " ".join(cur_val)

    for line in _sec_lines("caracteristicas"):
        ln = _n(line)
        if ln in PROP_MAP:
            _flush_cf()
            cur_key = PROP_MAP[ln]
            cur_val = []
        elif cur_key:
            cur_val.append(line)
    _flush_cf()

    # --- Beneficios ---
    beneficio_items: list[str] = []
    cur_bl = ""
    cur_bv: list[str] = []

    def _flush_b() -> None:
        if cur_bl:
            joined = " ".join(cur_bv)
            beneficio_items.append(f"{cur_bl}|{joined}" if joined else cur_bl)

    for line in _sec_lines("beneficios"):
        if ":" in line:
            idx = line.index(":")
            possible_lbl = line[:idx].strip()
            possible_val = line[idx + 1:].strip()
            if len(possible_lbl) < 60 and possible_lbl and possible_lbl[0].isupper():
                _flush_b()
                cur_bl = possible_lbl
                cur_bv = [possible_val] if possible_val else []
                continue
        if cur_bl:
            cur_bv.append(line)
        else:
            beneficio_items.append(line)
    _flush_b()
    propiedades_lista = "\n".join(beneficio_items)

    # --- Aplicaciones ---
    aplicaciones = "\n".join(l for l in _sec_lines("aplicaciones") if l != "•")

    # --- Modo de uso ---
    modo_uso = " ".join(l for l in _sec_lines("modo_uso") if l != "•")

    # --- Composición ---
    composicion: list[list[str]] = []
    for line in _sec_lines("composicion"):
        if _n(line).startswith("componente") or ("%" in line and _n(line).find("concentrac") != -1):
            continue
        parts = [p.strip() for p in line.split("|", 1)]
        if parts[0]:
            composicion.append([parts[0], parts[1] if len(parts) > 1 else ""])

    # --- Recomendaciones GHS ---
    GHS_CATS = {
        _n(k) for k in (
            "señal de peligro", "indicaciones h", "prevención", "prevención:",
            "respuesta", "almacenamiento", "eliminación", "eliminación:",
            "primeros auxilios", "epp requerido",
        )
    }
    ghs_items: list[str] = []
    cur_gc = ""
    cur_gv: list[str] = []

    def _flush_ghs() -> None:
        if cur_gc and cur_gv:
            ghs_items.append(f"{cur_gc}: {' '.join(cur_gv)}")

    for line in _sec_lines("ghs"):
        if line.startswith("Lote:"):
            break
        if ":" in line:
            idx = line.index(":")
            cat = line[:idx].strip()
            val = line[idx + 1:].strip()
            if _n(cat) in GHS_CATS:
                _flush_ghs()
                cur_gc = cat
                cur_gv = [val] if val else []
                continue
        if cur_gc:
            cur_gv.append(line)
    _flush_ghs()
    recomendaciones = "\n".join(ghs_items)

    # --- Lote ---
    lote = ""
    for line in lines:
        if line.startswith("Lote:"):
            lote = line[5:].strip()
            break

    return {
        "titulo": titulo,
        "nombre_producto": titulo.title() if titulo else "",
        "referencia": referencia,
        "sinonimos": sinonimos,
        "cas": cas,
        "fecha_revision": fecha_revision,
        "descripcion": descripcion,
        "caracteristicas_fisicas": cf,
        "propiedades_lista": propiedades_lista,
        "aplicaciones": aplicaciones,
        "modo_uso": modo_uso,
        "composicion": composicion,
        "recomendaciones": recomendaciones,
        "lote": lote,
    }


def ruta_archivo_biblioteca_segura(nombre: str) -> Path | None:
    """Ruta segura para FT simple o documento completo (FT+COA+SDS) en la biblioteca."""
    nombre = os.path.basename(nombre or "")
    if not nombre or nombre.startswith("~") or not _es_nombre_biblioteca(nombre):
        return None
    if _PREFIJO_COMPLETO.match(nombre):
        directorios = (COMPLETO_PDF_DIR,)
    else:
        directorios = (FICHAS_PDF_DIR, FICHAS_DIR)
    for directorio in directorios:
        if not directorio.exists():
            continue
        path = (directorio / nombre).resolve()
        try:
            path.relative_to(directorio.resolve())
        except ValueError:
            continue
        if path.is_file():
            return path
    return None
