#!/usr/bin/env python3
"""
Extrae fichas técnicas de los archivos Word en fichas_word/ y de las fichas
estructuradas en fichas_word/datos/*.yaml (formato FT/COA/SDS más reciente),
y genera PAGINA_WEB/site/data/fichas_tecnicas.json — el caché que lee el
"texto mágico" de Studio (app/tools/plantillas_texto_ia.py).

Estructura de salida:
{
  "ACIDO ASCORBICO": {
    "titulo": "ÁCIDO ASCÓRBICO",
    "descripcion": "...",
    "secciones": [{"titulo": "BENEFICIOS", "items": ["...", "..."]}, ...],
    "identidad": [["NOMBRE DEL PRODUCTO", "Ácido L Ascorbico"], ...],
    "propiedades": [["Apariencia", "Polvo blanco..."], ...],
    "microbiologia": [["Recuento total", "≤ 5000"], ...]
  }
}
"""

import os
import re
import json
from pathlib import Path
from docx import Document
import yaml

FICHAS_DIR = Path("/home/mckg/mi-agente/fichas_word")
DATOS_DIR  = FICHAS_DIR / "datos"
OUTPUT     = Path("/home/mckg/mi-agente/PAGINA_WEB/site/data/fichas_tecnicas.json")

# Cabeceras que indican inicio de tabla de identidad/propiedades
IDENTIDAD_KEYWORDS   = {"nombre del producto", "cas #", "cas#", "fórmula molecular",
                        "formula molecular", "sinónimos", "sinonimos", "nombre iupac"}
PROPIEDADES_KEYWORDS = {"apariencia", "peso molecular", "concentración", "concentracion",
                        "rotación", "perdidas", "pérdidas", "punto", "ph", "solubilidad",
                        "densidad", "viscosidad", "pureza", "humedad", "identificación"}
MICRO_KEYWORDS       = {"mesofilos", "hongos", "levaduras", "e. coli", "e. colí",
                        "salmonella", "recuento", "coliformes", "aerobios"}

SECTION_HEADINGS = {
    "descripción", "descripcion", "beneficios", "aplicaciones", "propiedades",
    "recomendaciones", "dosificación", "dosificacion", "estabilidad",
    "almacenamiento", "uso", "usos", "instrucciones", "advertencias",
    "propiedades físico-químicas", "propiedades microbiológicas",
    "propiedades fisico-quimicas", "propiedades microbiologicas",
}


def normalizar(texto: str) -> str:
    """Normaliza texto para comparación."""
    t = texto.lower().strip()
    # Eliminar acentos básicos
    for a, b in [("á","a"),("é","e"),("í","i"),("ó","o"),("ú","u"),("ñ","n")]:
        t = t.replace(a, b)
    # Eliminar caracteres no alfanuméricos excepto espacios
    t = re.sub(r"[^a-z0-9 ]", " ", t)
    t = re.sub(r"\s+", " ", t).strip()
    return t


def clave_desde_archivo(nombre_archivo: str) -> str:
    """FT ACIDO ASCORBICO.docx → ACIDO ASCORBICO"""
    nombre = nombre_archivo.replace(".docx", "").strip()
    if nombre.upper().startswith("FT "):
        nombre = nombre[3:].strip()
    return normalizar(nombre)


def clasificar_tabla(tabla) -> str:
    """Detecta si una tabla es de identidad, propiedades, o microbiología."""
    textos = []
    for row in tabla.rows[:3]:
        for cell in row.cells[:2]:
            textos.append(normalizar(cell.text))
    todos = " ".join(textos)
    hits_id   = sum(1 for k in IDENTIDAD_KEYWORDS   if k in todos)
    hits_prop = sum(1 for k in PROPIEDADES_KEYWORDS if k in todos)
    hits_micro= sum(1 for k in MICRO_KEYWORDS        if k in todos)
    if hits_micro >= 1: return "microbiologia"
    if hits_id   >= 1: return "identidad"
    if hits_prop >= 1: return "propiedades"
    return "otras"


def tabla_a_filas(tabla) -> list:
    """Convierte tabla a lista de [celda0, celda1]."""
    filas = []
    vistas = set()
    for row in tabla.rows:
        textos = [c.text.strip() for c in row.cells]
        # Eliminar celdas duplicadas (tablas con celdas unidas)
        clave = tuple(t[:30] for t in textos)
        if clave in vistas:
            continue
        vistas.add(clave)
        # Solo filas que tienen contenido
        contenido = [t for t in textos if t]
        if contenido and len(contenido) >= 2:
            filas.append([textos[0].strip(), textos[-1].strip()])
        elif contenido:
            filas.append([textos[0].strip(), ""])
    return filas


def extraer_ficha(path: Path) -> dict:
    """Lee un archivo .docx y extrae su contenido estructurado."""
    doc = Document(str(path))

    ficha = {
        "titulo": "",
        "descripcion": "",
        "secciones": [],          # [{"titulo": str, "items": [str]}]
        "identidad": [],          # [[label, valor], ...]
        "propiedades": [],
        "microbiologia": [],
    }

    # ── Extraer tablas por tipo ──────────────────────────────
    for tabla in doc.tables:
        tipo = clasificar_tabla(tabla)
        filas = tabla_a_filas(tabla)
        if tipo == "identidad":
            ficha["identidad"].extend(filas)
        elif tipo == "propiedades":
            ficha["propiedades"].extend(filas)
        elif tipo == "microbiologia":
            ficha["microbiologia"].extend(filas)

    # ── Extraer párrafos estructurados ──────────────────────
    seccion_actual = None
    items_actuales = []
    titulo_extraido = False

    for para in doc.paragraphs:
        texto = para.text.strip()
        if not texto:
            continue

        texto_norm = normalizar(texto)

        # Título principal (primer párrafo no vacío o estilo Normal corto)
        if not titulo_extraido and len(texto) < 80:
            ficha["titulo"] = texto
            titulo_extraido = True
            continue

        # ¿Es una cabecera de sección?
        es_cabecera = (
            para.style.name in ("Heading 1", "Heading 2", "Heading 3")
            or texto_norm in SECTION_HEADINGS
            or (len(texto) < 60 and texto.isupper())
        )

        if es_cabecera:
            # Guardar sección anterior
            if seccion_actual and items_actuales:
                ficha["secciones"].append({
                    "titulo": seccion_actual,
                    "items": items_actuales
                })
                items_actuales = []
            seccion_actual = texto
            # Si es DESCRIPCION, el siguiente contenido va ahí
        else:
            # Acumular items de la sección
            if seccion_actual:
                titulo_norm = normalizar(seccion_actual)
                if "descripci" in titulo_norm and not ficha["descripcion"]:
                    ficha["descripcion"] = texto
                else:
                    items_actuales.append(texto)

    # Guardar última sección
    if seccion_actual and items_actuales:
        ficha["secciones"].append({
            "titulo": seccion_actual,
            "items": items_actuales
        })

    # Si no se encontró descripción en párrafos, buscar en secciones
    if not ficha["descripcion"]:
        for sec in ficha["secciones"]:
            if "descripci" in normalizar(sec["titulo"]) and sec["items"]:
                ficha["descripcion"] = sec["items"][0]
                break

    return ficha


def clave_desde_nombre(nombre: str) -> str:
    return normalizar(nombre)


def extraer_ficha_yaml(path: Path) -> dict:
    """Lee una ficha estructurada fichas_word/datos/*.yaml y la convierte
    al mismo esquema que produce extraer_ficha() para los .docx."""
    with open(path, encoding="utf-8") as f:
        datos = yaml.safe_load(f) or {}

    ficha = {
        "titulo": (datos.get("titulo") or datos.get("nombre_producto") or "").strip(),
        "descripcion": (datos.get("descripcion") or "").strip(),
        "secciones": [],
        "identidad": datos.get("identidad") or [],
        "propiedades": datos.get("propiedades") or [],
        "microbiologia": (datos.get("_coa") or {}).get("parametros") or [],
    }

    beneficios = []
    for linea in datos.get("propiedades_lista") or []:
        titulo_prop, _, desc = str(linea).partition("|")
        titulo_prop = titulo_prop.strip()
        desc = desc.strip()
        beneficios.append(f"{titulo_prop}: {desc}" if desc else titulo_prop)
    if beneficios:
        ficha["secciones"].append({"titulo": "BENEFICIOS", "items": beneficios})

    aplicaciones = [str(a).strip() for a in (datos.get("aplicaciones") or []) if str(a).strip()]
    if aplicaciones:
        ficha["secciones"].append({"titulo": "APLICACIONES", "items": aplicaciones})

    modo_uso = (datos.get("modo_uso") or "").strip()
    if modo_uso:
        ficha["secciones"].append({"titulo": "MODO DE USO", "items": [modo_uso]})

    recomendaciones = (datos.get("recomendaciones") or "").strip()
    if recomendaciones:
        ficha["secciones"].append({"titulo": "RECOMENDACIONES", "items": [recomendaciones]})

    return ficha


def main():
    archivos = sorted(FICHAS_DIR.glob("FT *.docx"))
    print(f"Encontrados {len(archivos)} archivos FT *.docx")

    resultado = {}
    errores = []

    for archivo in archivos:
        clave = clave_desde_archivo(archivo.name)
        try:
            ficha = extraer_ficha(archivo)
            resultado[clave] = ficha
            print(f"  OK  {clave[:50]:<50}  secciones={len(ficha['secciones'])}  props={len(ficha['propiedades'])}")
        except Exception as e:
            errores.append((archivo.name, str(e)))
            print(f"  ERR {archivo.name}: {e}")

    yamls = sorted(
        p for p in DATOS_DIR.glob("*.yaml")
        if "plantilla_ejemplo" not in p.stem.lower()
    )
    print(f"\nEncontrados {len(yamls)} archivos datos/*.yaml")
    for archivo in yamls:
        try:
            ficha = extraer_ficha_yaml(archivo)
            nombre = ficha["titulo"] or archivo.stem
            clave = clave_desde_nombre(nombre)
            if not clave or clave == "nombre del producto":
                raise ValueError("sin titulo/nombre_producto real (plantilla sin completar)")
            score_nueva = (
                len(ficha["secciones"]) + len(ficha["propiedades"])
                + len(ficha["identidad"]) + bool(ficha["descripcion"])
            )
            existente = resultado.get(clave)
            if existente:
                score_existente = (
                    len(existente.get("secciones") or [])
                    + len(existente.get("propiedades") or [])
                    + len(existente.get("identidad") or [])
                    + bool(existente.get("descripcion"))
                )
                if score_nueva < score_existente:
                    print(f"  --  '{clave}' de {archivo.name} tiene menos datos que la ficha "
                          f"existente; se conserva la existente")
                    continue
                print(f"  (sobrescribe entrada existente '{clave}' con datos de {archivo.name})")
            resultado[clave] = ficha
            print(f"  OK  {clave[:50]:<50}  secciones={len(ficha['secciones'])}  props={len(ficha['propiedades'])}")
        except Exception as e:
            errores.append((archivo.name, str(e)))
            print(f"  ERR {archivo.name}: {e}")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT.write_text(json.dumps(resultado, ensure_ascii=False, indent=2))
    print(f"\n✓ Guardado: {OUTPUT}")
    print(f"  {len(resultado)} fichas  |  {len(errores)} errores")
    if errores:
        for n, e in errores:
            print(f"  ! {n}: {e}")


if __name__ == "__main__":
    main()
