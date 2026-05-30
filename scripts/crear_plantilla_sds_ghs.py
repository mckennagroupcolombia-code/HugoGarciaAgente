#!/usr/bin/env python3
"""
Genera plantilla SDS DOCX — formato GHS estilo Ventós (referencia: SDS ELEMI.pdf).
McKenna Group branding.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUT = Path(__file__).resolve().parents[1] / "fichas_word" / "plantillas" / "SDS PLANTILLA.docx"

# Títulos de sección como en SDS ELEMI (Ventós)
SECCIONES = [
    ("1", "IDENTIFICACIÓN DEL PRODUCTO", "{{S01}}"),
    ("2", "IDENTIFICACIÓN DEL PELIGRO O PELIGROS", "{{S02}}"),
    ("3", "COMPOSICIÓN/INFORMACIÓN SOBRE LOS COMPONENTES", "{{S03}}"),
    ("4", "PRIMEROS AUXILIOS", "{{S04}}"),
    ("5", "MEDIDAS DE LUCHA CONTRA INCENDIOS", "{{S05}}"),
    ("6", "MEDIDAS QUE DEBEN TOMARSE EN CASO DE VERTIDO ACCIDENTAL", "{{S06}}"),
    ("7", "MANIPULACIÓN Y ALMACENAMIENTO", "{{S07}}"),
    ("8", "CONTROLES DE LA EXPOSICIÓN/PROTECCIÓN PERSONAL", "{{S08}}"),
    ("9", "PROPIEDADES FÍSICAS Y QUÍMICAS", None),  # tabla dedicada
    ("10", "ESTABILIDAD Y REACTIVIDAD", "{{S10}}"),
    ("11", "INFORMACIÓN TOXICOLÓGICA", "{{S11}}"),
    ("12", "INFORMACIÓN ECOLÓGICA", "{{S12}}"),
    ("13", "INFORMACIÓN RELATIVA A LA ELIMINACIÓN DE LOS PRODUCTOS", "{{S13}}"),
    ("14", "INFORMACIÓN RELATIVA AL TRANSPORTE", "{{S14}}"),
    ("15", "INFORMACIÓN SOBRE LA REGLAMENTACIÓN", "{{S15}}"),
    ("16", "OTRAS INFORMACIONES", "{{S16}}"),
]

PROPIEDADES_FISICAS = [
    ("Aspecto", "{{PF_ASPECTO}}"),
    ("Color", "{{PF_COLOR}}"),
    ("Olor", "{{PF_OLOR}}"),
    ("Umbral olfativo", "{{PF_UMBRAL_OLOR}}"),
    ("pH", "{{PF_PH}}"),
    ("Punto de fusión/Punto de congelación", "{{PF_FUSION}}"),
    ("Punto de ebullición/rango de ebullición (ºC)", "{{PF_EBULICION}}"),
    ("Punto de Inflamación", "{{PF_INFLAMACION}}"),
    ("Velocidad de evaporación", "{{PF_EVAPORACION}}"),
    ("Inflamabilidad", "{{PF_INFLAMABILIDAD}}"),
    ("Límite inferior de inflamabilidad/explosividad", "{{PF_LIM_INF}}"),
    ("Límite superior de inflamabilidad/explosividad", "{{PF_LIM_SUP}}"),
    ("Presión de vapor", "{{PF_PRESION_VAPOR}}"),
    ("Densidad de vapor", "{{PF_DENS_VAPOR}}"),
    ("Densidad", "{{PF_DENSIDAD}}"),
    ("Densidad relativa", "{{PF_DENS_REL}}"),
    ("Solubilidad en agua", "{{PF_SOL_AGUA}}"),
    ("Solubilidad en otros disolventes", "{{PF_SOL_OTROS}}"),
    ("Coeficiente de partición n-octanol/agua", "{{PF_LOGP}}"),
    ("Temperatura de auto-ignición", "{{PF_AUTOIGNICION}}"),
    ("Temperatura de descomposición", "{{PF_DESCOMP}}"),
    ("Viscosidad, dinámica", "{{PF_VIS_DIN}}"),
    ("Viscosidad, cinemática", "{{PF_VIS_CIN}}"),
    ("Propiedades explosivas", "{{PF_EXPLOSIVAS}}"),
    ("Propiedades comburentes", "{{PF_COMBURENTES}}"),
]


def main() -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = Pt(50)
    sec.right_margin = Pt(50)

    # Encabezado estilo Ventós
    h = doc.add_paragraph("FICHA DE DATOS DE SEGURIDAD")
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.runs[0].bold = True
    h.runs[0].font.size = Pt(14)

    sub = doc.add_paragraph(
        "Conforme a los requerimientos del Sistema General Armonizado (GHS)"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    brand = doc.add_paragraph("M C K E N N A   G R O U P")
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand.runs[0].bold = True

    prod = doc.add_paragraph("{{TITULO}}")
    prod.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prod.runs[0].bold = True
    prod.runs[0].font.size = Pt(16)

    meta = doc.add_table(rows=2, cols=2)
    meta.rows[0].cells[0].text = "Fecha revisión: {{FECHA_REVISION}}"
    meta.rows[0].cells[1].text = "Versión: {{VERSION}}"
    meta.rows[1].cells[0].text = "Fecha impresión: {{FECHA_IMPRESION}}"
    meta.rows[1].cells[1].text = "Referencia: {{REFERENCIA}}"

    doc.add_paragraph()

    for num, titulo, placeholder in SECCIONES:
        sec_h = doc.add_paragraph(f"{num}. {titulo}")
        sec_h.runs[0].bold = True

        if num == "9":
            t = doc.add_table(rows=len(PROPIEDADES_FISICAS), cols=2)
            t.style = "Table Grid"
            for i, (label, ph) in enumerate(PROPIEDADES_FISICAS):
                t.rows[i].cells[0].text = label
                t.rows[i].cells[1].text = ph
        else:
            doc.add_paragraph(placeholder or "")

        doc.add_paragraph()

    foot = doc.add_paragraph(
        "MCKENNA GROUP S.A.S. · NIT 901316016-3 · Bogotá – Colombia · www.mckennagroup.co\n"
        "Formato SDS GHS (referencia Ventós/ELEMI) · Documento controlado."
    )
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"✓ Plantilla SDS Ventós/McKenna: {OUT}")


if __name__ == "__main__":
    main()
