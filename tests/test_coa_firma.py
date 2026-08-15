from pathlib import Path
import io
import shutil

from docx import Document
from PIL import Image

from PIL import ImageDraw

from app.services.coa import PLANTILLA_DEFAULT, aplicar_datos_a_docx
from app.services.coa_firma import (
    _parse_bbox,
    data_url_a_bytes,
    extraer_trazo_firma,
    recortar_firma_a_data_url,
)
from app.services.ficha_tecnica import _contexto_coa


def _datos_base() -> dict:
    return {
        "titulo": "PRODUCTO PRUEBA",
        "identificacion": {"nombre_comercial": "Producto prueba"},
        "lote": {"numero": "LOTE-1"},
        "parametros": [["Aspecto", "Polvo blanco", "Cumple"]],
        "empaque": {},
    }


def test_coa_sin_firmante_deja_solo_espacio(tmp_path: Path):
    salida = tmp_path / "coa-sin-firmante.docx"
    shutil.copy2(PLANTILLA_DEFAULT, salida)

    aplicar_datos_a_docx(salida, _datos_base())

    texto = Document(str(salida)).tables[13].rows[0].cells[0].text
    assert "REVISADO Y APROBADO POR" in texto
    assert "______________________________" in texto
    assert "Armando García Velandia" not in texto
    assert "Asesor de Calidad" not in texto


def test_coa_imprime_datos_de_firma_extraidos(tmp_path: Path):
    datos = _datos_base()
    datos["firma"] = {
        "nombre": "Laura Pérez",
        "cargo": "Quality Manager",
        "organizacion": "Laboratorio Ejemplo",
    }
    salida = tmp_path / "coa-con-firmante.docx"
    shutil.copy2(PLANTILLA_DEFAULT, salida)

    aplicar_datos_a_docx(salida, datos)

    texto = Document(str(salida)).tables[13].rows[0].cells[0].text
    assert "Laura Pérez" in texto
    assert "Quality Manager" in texto
    assert "Laboratorio Ejemplo" in texto

    contexto = _contexto_coa(datos)
    assert contexto["firma_nombre"] == "Laura Pérez"
    assert contexto["firma_cargo"] == "Quality Manager"
    assert contexto["firma_organizacion"] == "Laboratorio Ejemplo"


def test_parse_bbox_0_1000():
    assert _parse_bbox([700, 100, 850, 450]) == (0.7, 0.1, 0.85, 0.45)


def test_recorta_linea_de_firma_a_data_url():
    img = Image.new("RGB", (200, 200), (255, 255, 255))
    for x in range(20, 90):
        for y in range(140, 160):
            img.putpixel((x, y), (10, 10, 10))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    data_url = recortar_firma_a_data_url(buf.getvalue(), "image/png", [700, 100, 850, 450])
    assert data_url is not None
    assert data_url.startswith("data:image/png;base64,")

    datos = _datos_base()
    datos["firma"] = {"imagen_b64": data_url, "nombre": "Ana Ruiz"}
    ctx = _contexto_coa(datos)
    assert ctx["firma_imagen_src"].startswith("data:image/png;base64,")
    assert ctx["firma_nombre"] == "Ana Ruiz"


def _firma_sobre(fondo: tuple[int, int, int], tinta: tuple[int, int, int]) -> bytes:
    img = Image.new("RGB", (240, 90), fondo)
    trazo = ImageDraw.Draw(img)
    trazo.line([(20, 60), (60, 25), (100, 65), (150, 30), (210, 55)], fill=tinta, width=3)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _abrir_rgba(data_url: str) -> Image.Image:
    decodificado = data_url_a_bytes(data_url)
    assert decodificado is not None
    return Image.open(io.BytesIO(decodificado[0]))


def test_extrae_trazo_sobre_papel_gris_deja_fondo_transparente():
    """El recorte del escáner suele traer papel gris, no blanco puro."""
    data_url = extraer_trazo_firma(_firma_sobre((178, 176, 172), (45, 45, 50)))
    assert data_url is not None

    res = _abrir_rgba(data_url)
    assert res.mode == "RGBA"
    alpha = list(res.getchannel("A").get_flattened_data())
    transparentes = sum(1 for v in alpha if v == 0)
    assert res.getchannel("A").getpixel((0, 0)) == 0
    assert max(alpha) == 255, "el trazo debe quedar opaco"
    assert transparentes > len(alpha) * 0.7, "el fondo del papel debe desaparecer"


def test_extrae_trazo_recorta_al_contenido():
    data_url = extraer_trazo_firma(_firma_sobre((255, 255, 255), (10, 10, 10)))
    assert data_url is not None
    res = _abrir_rgba(data_url)
    assert res.size[0] < 240 and res.size[1] < 90


def test_extrae_trazo_ignora_recorte_sin_contraste():
    buf = io.BytesIO()
    Image.new("RGB", (120, 60), (255, 255, 255)).save(buf, format="PNG")
    assert extraer_trazo_firma(buf.getvalue()) is None


def test_recorte_de_firma_llega_sin_fondo_al_pdf():
    """La firma escaneada debe entrar al COA ya recortada y transparente."""
    img = Image.new("RGB", (200, 200), (250, 250, 250))
    for x in range(30, 170):
        for y in range(120, 180):
            img.putpixel((x, y), (185, 183, 180))  # papel fotografiado
    trazo = ImageDraw.Draw(img)
    trazo.line([(45, 160), (80, 135), (120, 165), (160, 138)], fill=(35, 35, 40), width=3)
    buf = io.BytesIO()
    img.save(buf, format="PNG")

    data_url = recortar_firma_a_data_url(buf.getvalue(), "image/png", [600, 150, 900, 850])
    assert data_url is not None
    res = _abrir_rgba(data_url)
    assert res.mode == "RGBA"
    assert res.getchannel("A").getpixel((0, 0)) == 0
